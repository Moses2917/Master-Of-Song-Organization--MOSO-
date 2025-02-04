from time import time as today
from json import load
from os import environ as ENV
from docx import Document
from regex import D
basePth = ENV.get("OneDrive")
#open all files, get latest version and store lyrics in dict
# if songLyrics.get('latestChange',None) or do < today()

#templates for songLyrics
def template(songs_to_open_bool=True) -> dict:
    songLyrics = {
        'old': {},
        'new': {},
        'latestChange': today()
    }
    songs_to_open = {
        'old': {},
        'new': {}
    }
    return songs_to_open if songs_to_open_bool else songLyrics

def getAllLyricsDict() -> dict:
    with open('AllLyrics.json', 'r', encoding='utf-8') as f:
        return load(f)       

def GetAllSongPths(songs_to_open,index:str, book = 'old'or'new',) -> dict:
    try:
        for songNum in index["SongNum"]:
            songPth_realative = index["SongNum"][songNum]['latestVersion']
            songPth_full = f"{basePth}\\{songPth_realative}"
            songs_to_open[book][songNum] = songPth_full
    except:
        print(f'This song threw an error: {index["SongNum"][songNum]}')
    return songs_to_open

def readLyrics(filePth:str) -> str:
    lyrics = ''
    doc = Document(filePth)
    for p in doc.paragraphs:
        lyrics += p.text
    
    if lyrics != '':
        return lyrics
    return None

def readLyrics(doc:Document) -> str:
    lyrics = ''
    # doc = Document(filePth)
    for p in doc.paragraphs:
        lyrics += p.text
    
    if lyrics != '':
        return lyrics
    return None

def updateSongLyrics(book:str, songNum:str, lyrics:Document):
    allLyrics = getAllLyricsDict()
    if book.lower() == 'old':
        allLyrics['old'][songNum] = readLyrics(lyrics)
        with open('AllLyrics.json', 'w', encoding='utf-8') as f:
            from json import dump
            dump(allLyrics,f,ensure_ascii=False,indent=4)
        return True
    elif book.lower() == 'new':
        allLyrics['new'][songNum] = readLyrics(lyrics)
        with open('AllLyrics.json', 'w', encoding='utf-8') as f:
            from json import dump
            dump(allLyrics,f,ensure_ascii=False,indent=4)
        return True
    return None

def updateAllLyrics():
    #Gather all current songs.
    with open('AllLyrics.json', 'r', encoding='utf-8') as f:
        allLyrics = load(f)
    
    from os import environ as ENV

    onedrive = ENV.get('onedrive')

    from os import path

    

    return True

def getAllLyrics():
    songs_to_open = template()
    songLyrics = template(songs_to_open_bool=False)
    #Start from olds
    with open('wordSongsIndex.json', 'r', encoding='utf-8') as f:
        oldSongs = load(f)
        GetAllSongPths(oldSongs,'old')
    # Close file and continue with reading all red songs
    with open('REDergaran.json', 'r', encoding='utf-8') as f:
        newSongs = load(f)
        GetAllSongPths(newSongs,'new')
    # with open('test.json', 'w', encoding='utf-8') as f:
    #     from json import dump
    #     dump(songs_to_open,f,ensure_ascii=False,indent=4)
    for book in songs_to_open:
        for songNum in songs_to_open[book]:
            filePth = songs_to_open[book][songNum]
            lyrics = readLyrics(filePth)
            songLyrics[book][songNum] = lyrics
    with open('AllLyrics.json', 'w', encoding='utf-8') as f:
        from json import dump
        dump(songLyrics,f,ensure_ascii=False,indent=4)

def WordToJson(songNum, book):
    """
    PORTED FROM APP.PY
    Opens a word document based on the provided song number and book, 
    and returns the contents of the document as a JSON object.

    Args:
        songNum (int): The number of the song to be opened.
        book (str): The name of the book containing the song.

    Returns:
        dict: A JSON formatted string containing the contents of the word document.
    """
    from docx import Document
    #Used to find and open word doc, sends back a html formated str of it

    with open("REDergaran.json", 'r', encoding='utf-8') as f:
        index:dict = load(f)
    
    if index["SongNum"].get(songNum,None):
        songPth = index["SongNum"][songNum]["latestVersion"]
        filePth = ENV.get("OneDrive")+"\\"+songPth #find pth from index, and attach the location for onedrive
        doc = Document(filePth) #load doc file
        docParagraphs = doc.paragraphs # returns a list of doc paragrpahs from which text will be extracted
        text = ''
        
        for para in docParagraphs:
            text += para.text + '\n'
        
        # Split the text into chunks based on line breaks
        chunks = text.split('\n\n')
        
        # Convert chunks to HTML
        html_chunks = []
        for chunk in chunks:
            lines = chunk.split('\n')
            html_lines = ['<p>' + line + '</p>' for line in lines]
            html_chunk = ''.join(html_lines)
            html_chunks.append(html_chunk)

        # Join the chunks with line breaks, adding or subtracting br will add or subtract the breaks between the paragraphs
        html_text = '<br>'.join(html_chunks)
        return html_text

def wordsToJson():
    with open("REDergaran.json", 'r', encoding="utf-8") as f:
        Songs = load(f)
    songLyrics = {}
    for song in Songs["SongNum"]:
        if song != "SongNum":
            fp = ENV.get("OneDrive") + "/" + Songs["SongNum"][song]["latestVersion"]
            # print(fp)
            doc = Document(fp) #load doc file
            docParagraphs = doc.paragraphs # returns a list of doc paragrpahs from which text will be extracted
            text = ''
            
            for para in docParagraphs:
                text += para.text + '\n'
            
            songLyrics[song] = text
    
    with open('AppLyrics.json', 'w', encoding='utf-8') as f:
        from json import dump
        dump(songLyrics,f,ensure_ascii=False,indent=4)

def singleWordToJson(songNum:str):
    with open("REDergaran.json", 'r', encoding="utf-8") as f:
        Songs = load(f)
    
    #load the dir to update
    with open('AppLyrics.json', 'r', encoding='utf-8') as f:
        songLyrics = load(f)
        fp = ENV.get("OneDrive") + "/" + Songs["SongNum"][songNum]["latestVersion"]
        # print(fp)
        doc = Document(fp) #load doc file
        docParagraphs = doc.paragraphs # returns a list of doc paragrpahs from which text will be extracted
        text = ''
        
        for para in docParagraphs:
            text += para.text + '\n'
        
        songLyrics[songNum] = text
    
    with open('AppLyrics.json', 'w', encoding='utf-8') as f:
        from json import dump
        dump(songLyrics,f,ensure_ascii=False,indent=4)

if __name__ == "__main__":
    pass
    singleWordToJson('95')
    singleWordToJson('96')
    # onedrive = ENV.get('onedrive')
    # updateSongLyrics('old','389',Document(onedrive+'\\Word songs/389 Տոն է այսոր սուրբ հաղթական.docx'))
