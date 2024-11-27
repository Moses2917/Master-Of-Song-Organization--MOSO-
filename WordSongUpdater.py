import re, time, docx, json
from os import path as pth, remove, environ
from docx.shared import Pt

from getAllLyrics import updateSongLyrics

month = time.strftime('%m')
year = time.strftime('%y')
fullYear = time.strftime('%Y')
day = time.strftime('%d')

oldBook_pth = "wordSongsIndex.json"
Ergaran_pth = "REDergaran.json" # after combining red and ergaran this is now the main index file for the red book songs
redErgaran_pth = "oldpythfiles\REDergaran.json" #However for extra redunency I will leave this in

# def sortIndex(index):
#     """sorts a json index

#     Args:
#         index (json): a json obj, is song index
#     """    
#     ergaran["SongNum"] = {}
#     for item in sorted_items:
#         ergaran["SongNum"][item[0]] = item[1]
#     #returns a sorted index
#     return dict(sorted(index["SongNum"].items(), key=lambda x: int(x[0])))
# with open(oldBook_pth,"r",encoding="utf-8") as f:
#     sortedIndex = json.load(f)
#     sortedIndex["SongNum"] = sortIndex(sortedIndex["SongNum"])
# with open(oldBook_pth, "w", encoding='utf-8') as f:
#     json.dump(sortedIndex,f,indent=4,ensure_ascii=False)

def getDocText(filename):
    text = ""

    doc = docx.Document(filename)
    for p in doc.paragraphs:
        text += p.text + "\n"
    return text

def getOldSongTitle(oldSong_text):
    tab = re.findall("\n.+", oldSong_text) # for finding title
    # print(tab[1])
    title = tab[1]
    # Use re.sub() to remove "1." and "," from the text
    title = re.sub(r"1.", "", title)
    title = re.sub(r",", "", title)
    title = re.sub(r"\n", "", title)
    print(title)
    return title

#gets first line of song and uses that as title
def getRedSongTitle(text):
    return re.findall("\n(\d.*)\n",text)[0] 

#Finds the biggest version and returns it as an int
def latestVer(jsonIndex, songNum:str):
    """Finds latest version of given songNum in given json index & returns it,
    if it can't find it creates one with empty Title, Version/Latest Version"""
    latestV = 0
    try:
        for attr in jsonIndex["SongNum"][songNum]:
            if "v" in attr:
                v = int(re.findall(r"([\d].*)", attr)[0])
                if v > latestV:
                    latestV = v            
        print("The latest version is:", latestV)
        if latestV > 4: #limits the version amounts to 5, until I can do something about them
            latestV = 4
        return latestV, jsonIndex["SongNum"][songNum]['Title']
    except:
        print("This song number has not been indexed yet")
        print("Adding index...")
        jsonIndex['SongNum'][songNum] = {} #create a directory with that song number, and empty values
        jsonIndex['SongNum'][songNum]["Title"] = ""        
        jsonIndex['SongNum'][songNum]["latestVersion"] = ""
        jsonIndex['SongNum'][songNum]["v1"] = latestV
        return latestV, jsonIndex["SongNum"][songNum]['Title']
#had to make a seperate functions bc python does not support function overloading :(
def latestVerErg(jsonIndex, songNum:str):
    """Finds latest version of given songNum & the song title in given json index & returns it,
    if it can't find it creates one with empty Title, Version/Latest Version"""
    latestV = 0
    try:
        for attr in jsonIndex["SongNum"][songNum]:
            if "v" in attr:
                v = int(re.findall(r"([\d].*)", attr)[0])
                if v > latestV:
                    latestV = v
        
        print("The latest version is:", latestV)
        return latestV, jsonIndex["SongNum"][songNum]['Title']
    except:
        print("This song number has not been indexed yet")
        print("Loading another index...")
        with open(redErgaran_pth, "r", encoding='utf-8') as f:
            redErgaran_Index = json.load(f)
            # latestVerErg(redErgaran_Index)   #lol imagine making this recursive
        if songNum in redErgaran_Index['SongNum']:
            jsonIndex['SongNum'][songNum] = redErgaran_Index['SongNum'][songNum] #create a directory with that song number, and values from REDergaran.json
            
            #This for loop is a bit redundent, however I want to cover 100% of all cases, 
            for attr in redErgaran_Index["SongNum"][songNum]:
                if "v" in attr:
                    v = int(re.findall(r"([\d].*)", attr)[0])
                    if v > latestV:
                        latestV = v            
            print("The latest version is:", latestV)
            if latestV > 4: #limits the version amounts to 5, until I can do something about them
                latestV = 4
            return latestV, redErgaran_Index["SongNum"][songNum]['Title']
        else:
            jsonIndex['SongNum'][songNum] = {} #create a directory with that song number, and empty values
            jsonIndex['SongNum'][songNum]["Title"] = ""
            jsonIndex['SongNum'][songNum]["latestVersion"] = ""
            jsonIndex['SongNum'][songNum]["v1"] = latestV
            return latestV, jsonIndex["SongNum"][songNum]['Title']

def getDocTextAndIndentation(filename:str):
    """
    Reads a Word document file and extracts the text and indentation information 
    from each paragraph. It identifies song sections based on specific start and 
    end indicators in the text and processes the paragraphs accordingly. Then it saves the file.

    Parameters:
        filename (str): The path to the Word document file to be read.

    Returns:
        A list of dictionaries containing the text and indentation information 
        for each paragraph in the song sections.
    """
    
    doc = docx.Document(filename)
    

    text_and_indentation = [] #turn into a list of dicts
    song = []
    bookOld = False
    first = True
    for p in doc.paragraphs:
        
        if "[start:song" in p.text:
            my_doc = docx.Document()
            song = []
            songNum = None
            first = True
            if "old" in p.text: #Possible starting loc, or just make the doc file in it's entirety and and send off a list of docs to be saved somewhere else
                bookOld = True
                
            #have to add bc the songNum gets shoved in with the start indicator sometimes: '[start:song]\n171'
            if (re.search(r"[0-9]",p.text)):
                songNum = re.sub(r"\D", "", p.text)
                first = False
                
        if not("end" in p.text or "start" in p.text):
            if first: 
                songNum = p.text.split("\n")[0]
                first=False
            first_line_indent = p.paragraph_format.first_line_indent
            left_indent = p.paragraph_format.left_indent
            right_indent = p.paragraph_format.right_indent

            Placeholder = my_doc.add_paragraph(p.text)
            Placeholder.paragraph_format.space_after = 0
            if first_line_indent is not None:
                Placeholder.paragraph_format.first_line_indent = first_line_indent
            if left_indent is not None:
                Placeholder.paragraph_format.left_indent = left_indent
            if right_indent is not None:
                Placeholder.paragraph_format.right_indent = right_indent

            song.append({
                'text': p.text,
                # 'book': re.findall(pattern, p.text,re.DOTALL)[0],
                # 'old': bookOld,
                'first_line_indent': first_line_indent,
                'left_indent': left_indent,
                'right_indent': right_indent
            })
        if "end" in p.text: # Def ending loc
            if doc.paragraphs[1].text:
                saveDocFromDoc(my_doc, bookOld, songNum)
            else:
                print("No Pass!")
            #push song to text var and reset song var
            bookOld = False
            song = []


def saveDocFromDoc(song_Doc: docx, oldBook:str, songNum:str):
    # method logic:
    # 1. check if song exists in relevant book/index
    # 2. if so, append to index, "version", and "latestVersion"
    # 2.1 Based on this info the program needs to generate a suitable filepath to be logged and saved.
    # 2.5 if not, append to index, "SongNum", "Title", "version", and "latestVersion"
    # 3. once the lv(LatestVersion) is acquired, then add one and get new var cv(Current Version)
    # 4. Given the cv now create file path and amend to json index
    # 4.5 Save at the newly created file path
    # 4.7 If tues/thurs no need to save

    #Possibly add a check where it checks to see if any songs have the same title and updates with the latest ie: song '3' and '3 ' should be combined so as to not cause later confusion

    if songNum == None:
        return
    
    if oldBook:
        with open(oldBook_pth, 'r', encoding='utf-8') as f:
            oldBook_Index = json.load(f)

        # filepth = "relative filepath" + latestVer(jsonIndex=oldBook_pth, songNum=songNum) # type: ignore
        cv, title = latestVer(jsonIndex=oldBook_Index, songNum=songNum)#add 1 to get the current version
        cv += 1
        print(oldBook_Index["SongNum"][songNum])
        # print("Debug String..")

        base_file_path = "Word songs/{} {} v{}.docx".format(str(songNum), title.split("\n")[0], str(cv))
        # base_file_path = base_file_path.split("\n")[0] #Get issues when saving multiple lines ex: 'Word songs/389 Տոն է այսոր սուրբ հաղթական\nՀիսուս կուգա Երուսաղեմ v2.docx'
        oldBook_Index["SongNum"][songNum]["v"+ str(cv)] = base_file_path
        oldBook_Index["SongNum"][songNum]["latestVersion"] = base_file_path
        style = song_Doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(22)
        song_Doc.save("C:/Users/{}/OneDrive/".format(environ.get("USERNAME")) + base_file_path)
        with open(oldBook_pth, 'w', encoding='utf-8') as f:
            #oldBook_Index["SongNum"] = dict(sorted(oldBook_Index["SongNum"].items(), key=lambda x: int(x[0]))) # should sort the songs before saving
            json.dump(oldBook_Index, f, indent=4, ensure_ascii=False)
        print(base_file_path)

    if not oldBook:
        # compare songNum with ergaran.json index, if not there pull info from REDergarn.json
        with open(Ergaran_pth, "r", encoding='utf-8') as f:
            Book_Index = json.load(f)

        # filepth = "relative filepath" + latestVer(jsonIndex=oldBook_pth, songNum=songNum) # type: ignore
        cv, title = latestVerErg(jsonIndex=Book_Index, songNum=songNum)#add 1 to get the current version
        cv += 1
        print(Book_Index["SongNum"][songNum])
        # print("Debug String..")  # Note: Funny enough the test num I used does not have a corresponding title, which does not really matter that much, however I could add some functionality to fill it later on
        # However I'm not so sure about just saving files in ergaran as songnum.docx like I already do in red ergaran

        base_file_path = "Երգարան Word Files/{} {} v{}.docx".format(str(songNum), title.split("\n")[0], str(cv))
        Book_Index["SongNum"][songNum]["v"+ str(cv)] = base_file_path
        Book_Index["SongNum"][songNum]["latestVersion"] = base_file_path
        style = song_Doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(22)        
        song_Doc.save("C:/Users/{}/OneDrive/".format(environ.get("USERNAME")) + base_file_path)
        with open(Ergaran_pth, 'w', encoding='utf-8') as f:
            #Book_Index["SongNum"] = dict(sorted(Book_Index["SongNum"].items(), key=lambda x: int(x[0]))) # should sort the songs before saving
            json.dump(Book_Index, f, indent=4, ensure_ascii=False)
        print(base_file_path)
    #TODO: make a function to update lyrics
    if not oldBook:
        if updateSongLyrics(book="new", songNum=songNum, lyrics = song_Doc) == None:
            print("Lyrics not Updated")
            print(f"Song Num: {songNum} Book: New")
        else:
            print("Lyrics Updated")
    else:
        if updateSongLyrics(book="old", songNum=songNum, lyrics = song_Doc) == None:
            print("Lyrics not Updated")
            print(f"Song Num: {songNum} Book: New")
        else:
            print("Lyrics Updated")

def getNums(filename: str):
    """Reads the file and returns a dict with the text along with a bool if it is from the old book"""
    doc = docx.Document(filename)
    SongList = []
    bookOld = False
    first = True
    songNum = None
    for p in doc.paragraphs:
        if "[start:song" in p.text:
            songNum = None
            first = True
            if "old" in p.text: #Possible starting loc, or just make the doc file in it's entirety and and send off a list of docs to be saved somewhere else
                bookOld = True
                
            #have to add bc the songNum gets shoved in with the start indicator sometimes: '[start:song]\n171'
            if (re.search(r"[0-9]",p.text)):
                songNum = re.sub(r"\D", "", p.text)
                first = False
                
        if not("end" in p.text or "start" in p.text):
            if first: 
                songNum = p.text.split("\n")[0]
                first=False

        if "end" in p.text:  # Def ending loc
            if bookOld == False: bookOld = "New"
            else: bookOld = "Old"
            # print(bookOld, songNum)
            SongList.append((bookOld, songNum))
            bookOld = False
            
    # return text_and_indentation
    return str(SongList)
    # return  SongList

# print(getDocTextAndIndentation(r"C:\Users\Armne\OneDrive\Երգեր\03.2024\03.17.24.docx")) #am getting weird error