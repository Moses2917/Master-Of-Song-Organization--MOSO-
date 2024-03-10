import os
import json
import re
from docx import Document


def extract_song_info(file_path):
    doc = Document(file_path)
    song_number = re.search(r'\d+', os.path.basename(file_path)).group()

    # Extract the first paragraph text
    first_paragraph_text = doc.paragraphs[0].text.strip()

    # Remove any leading numbers, periods, and whitespace from the first paragraph text
    cleaned_first_paragraph_text = re.sub(r'^\d+\.?\s*', '', first_paragraph_text)

    # Check if the cleaned first paragraph text is not empty and doesn't start with '-'
    if cleaned_first_paragraph_text and not cleaned_first_paragraph_text.startswith('-'):
        title = cleaned_first_paragraph_text
    else:
        # If the first paragraph doesn't seem to contain the title, search for the title in subsequent paragraphs
        title = None
        for paragraph in doc.paragraphs[1:]:
            paragraph_text = paragraph.text.strip()
            cleaned_paragraph_text = re.sub(r'^\d+\.?\s*', '', paragraph_text)
            if cleaned_paragraph_text and not cleaned_paragraph_text.startswith('-'):
                title = cleaned_paragraph_text
                break
    file_path = file_path.split('\\Word songs\\pptSong')
    file_path = 'Word songs\\pptSong' + file_path[1]
    return song_number, title, file_path

def main():
    word_folder = '{}\\Word songs\\pptSong'.format(os.environ.get("ONEDRIVE"))
    song_info = {}

    for root, dirs, files in os.walk(word_folder):
        for filename in files:
            if filename.endswith('.docx'):
                file_path = os.path.join(root, filename)
                song_number, title, file_path = extract_song_info(file_path)
                if song_number and title:
                    song_data = {
                        'Title': title,
                        'latestVersion': file_path,
                        'v1': file_path
                    }
                    song_info[song_number] = song_data

    with open('song_info.json', 'w', encoding='utf-8') as json_file:
        json.dump(song_info, json_file, ensure_ascii=False, indent=4)

if __name__ == '__main__':
    main()