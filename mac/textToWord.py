import time
import docx
from docx.shared import Pt
from glob import glob
## Stuff for doc2txt
import xml.etree.ElementTree as ET
from docx.shared import Pt
import os
## use a try method to iterate through already made word files, and just in case go to the red words folder
my_doc = docx.Document("Users\movsesmovsesyan\Library\CloudStorage\OneDrive-Personal\Choir Songs Template.docx")

time.strftime('%B')

wordFile = open("Choir Songs Template.dotx", 'r', encoding='utf_8')
template = wordFile.read()
wordFile.close
month = time.strftime('%m')
year = time.strftime('%y')
fullYear = time.strftime('%Y')
day = time.strftime('%d')
# word = open("C:/Users/moses/OneDrive/Երգեր/" + month + "." + fullYear + "/" + month + "." + day + ".P" + year + ".docx", 'a', encoding='utf_8')
# word.write(template)
songs = []
imp = []

print("Press any key except for 'n' or 'o' to cancel the loop")
while True:
    inp = input("Enter 'n' or 'o' ")
    if inp == 'o' or inp == 'n':
        imp.append(inp)
        s = input("song num: ")
        songs.append(s)
        ##make a list/ array to store the values of which song num is old or new
    else:
        break



def xml2text(xml):
    """
    A string representing the textual content of this run, with content
    child elements like ``<w:tab/>`` translated to their Python
    equivalent.
    Adapted from: https://github.com/python-openxml/python-docx/
    """
    text = u''
    root = ET.fromstring(xml)
    for child in root.iter():
        if child.tag == qn('w:t'):
            t_text = child.text
            text += t_text if t_text is not None else ''
        elif child.tag == qn('w:tab'):
            text += '\t'
        elif child.tag in (qn('w:br'), qn('w:cr')):
            text += '\n'
        elif child.tag == qn("w:p"):
            text += '\n\n'
    return text

def process(filename):
    text = ""
    doc = docx.Document(filename)
    for p in doc.paragraphs:
        # if doc.text #if doc.text has a tab(\t) then add a tab stop
        text += p.text + "\n"
    return text


# songs: 350 251
# imp:    o  n
for x in songs:
    y = songs.index(x)
    if 'n' in imp[y]:
        print(x + "\n")
        try:
            filename = glob("C:/Users/moses/OneDrive/Երգարան Word Files/" + x + "*.docx")[0] #gets name
            doc = docx.Document(filename)
            # style = doc.styles
            tempFile = process(filename)
        except:
            songFile = open("C:/Users/moses/OneDrive/Ergaran Web Word Songs/" + x + ".txt", 'r', encoding='utf_8')
            tempFile = songFile.read()
            songFile.close()
    else:
        print(x + "\n")
        try:
            filename = glob("C:/Users/Armne/OneDrive/Word songs/" + x + "*")[0] # gets name
            doc = docx.Document(filename)
            tempFile = process(filename)
        except:
            print(x + ":" + " FileNotFoundError")
            tempFile = ""
        
    # word.write("\n" + tempFile + '\n')
    my_doc.add_paragraph(tempFile + '\n')#.add_run()


print("Adjusting for readability....")
font = my_doc.styles['Normal'].font
font.name = 'Arial'
font.size = Pt(22)
print("it is done")
# word.close()
# my_doc.save("C:/Users/moses/OneDrive/Երգեր/" + month + "." + fullYear + "/" + month + "." + day + "." + year + "TEST FILE DO NOT CLICK" + ".docx")
my_doc.save("C:/Users/moses/OneDrive/Երգեր/" + month + "." + fullYear + "/" + month + "." + day + "." + year + "PRACTICETEST.docx")

