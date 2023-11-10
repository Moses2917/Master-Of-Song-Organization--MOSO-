import time
import docx
from docx.shared import Pt
from glob import glob
## Stuff for doc2txt
import xml.etree.ElementTree as ET
import os
## use a try method to iterate through already made word files, and just in case go to the red words folder
my_doc = docx.Document("C:/Users/Armne/OneDrive/Choir Songs Template.docx")

time.strftime('%B')

wordFile = open("Choir Songs Template.dotx", 'r', encoding='utf_8')
template = wordFile.read()
wordFile.close
month = time.strftime('%m')
year = time.strftime('%y')
fullYear = time.strftime('%Y')
day = time.strftime('%d')
# word = open("C:/Users/moses/OneDrive/Երգեր/" + month + "." + fullYear + "/" + month + "." + day + ".P" + year + ".docx", 'a', encoding='utf_8')
# word.write(template)
songs = []
imp = []

book = input("Are all the songs from new/red songs\n y|n\n")
if("y" == book):
    while True:##Assumes that all songs are from new.
        imp.append("n")
        s = input("song num: ")
        if s == "":
            break
        else:
            songs.append(s)
else:
    print("Press any key except for 'n' or 'o' to cancel the loop")
    while True:
        inp = input("Enter 'n' or 'o' ")
        if inp == 'o' or inp == 'n':
            imp.append(inp)
            s = input("song num: ")
            songs.append(s)
            ##make a list/ array to store the values of which song num is old or new
        else:
            break

def process(filename):
    text = ""
    doc = docx.Document(filename)
    for p in doc.paragraphs:
        # if doc.text #if doc.text has a tab(\t) then add a tab stop
        text += p.text + "\n"
    return text

# Major bug, 16 if not found the module automatically goes to 164, which is a big problem
# songs: 350 251
# imp:    o  n
def oldAndNew(doc):
    for x in songs:
        y = songs.index(x)
        if 'n' in imp[y]:
            print(x + "\n")
            try:
                filename = glob("C:/Users/Armne/OneDrive/Երգարան Word Files/" + x + "*.docx")[0] #gets name

                doc = docx.Document(filename)
                # style = doc.styles
                tempFile = process(filename)
            except:
                songFile = open("C:/Users/Armne/OneDrive/RED Words/" + x + ".txt", 'r', encoding='utf_8')
                tempFile = songFile.read()
                songFile.close()
        else:
            print(x + "\n")
            try:
                filename = glob("C:/Users/Armne/OneDrive/Word songs/" + x + "*")[0] # gets name
                doc = docx.Document(filename)
                tempFile = process(filename)
            except:
                print(x + ":" + " FileNotFoundError")
                tempFile = ""
            
        # word.write("\n" + tempFile + '\n')
        my_doc.add_paragraph(tempFile + '\n')#.add_run()
    return my_doc


def songCollector(doc):
    #empty for now
    print(None)

for x in songs:
    y = songs.index(x)
    if 'n' in imp[y]:
        print(x + "\n")
        try:
            filename = glob("C:/Users/Armne/OneDrive/Երգարան Word Files/" + x + "*.docx")[0] #gets name
            tempFile = "[start:song]\n" + process(filename) + "[end:song]"
        except:
            tempFile = "[start:song]\n" + process("C:/Users/Armne/OneDrive/RED Words/" + str(x) + ".docx") + "[end:song]"
    else:
        print(x + "\n")
        try:
            filename = glob("C:/Users/Armne/OneDrive/Word songs/" + x + "*")[0] # gets name
            doc = docx.Document(filename)
            tempFile = "[start:song:old]\n" + process(filename) + "[end:song:old]"
        except:
            print(x + ":" + " FileNotFoundError")
            tempFile = ""
        
    # word.write("\n" + tempFile + '\n')
    my_doc.add_paragraph(tempFile + '\n')#.add_run()


print("Adjusting for readability....")
font = my_doc.styles['Normal'].font
font.name = 'Arial'
font.size = Pt(22)
print("it is done")
# word.close()
# my_doc.save("C:/Users/moses/OneDrive/Երգեր/" + month + "." + fullYear + "/" + month + "." + day + "." + year + "TEST FILE DO NOT CLICK" + ".docx")
my_doc.save("C:/Users/Armne/OneDrive/Երգեր/" + month + "." + fullYear + "/" + month + ".PORC " + day + "." + year + "PORC.docx")


# for x in songs:
#     print(x + "\n")
#     try:
#         filename = glob("C:/Users/Armne/OneDrive/Word songs/" + x + "*")[0] # gets name
#         doc = docx.Document(filename)
#         tempFile = process(filename)
#     except:
#         print(x + ":" + " FileNotFoundError")
#         tempFile = ""
        
#     # word.write("\n" + tempFile + '\n')
#     my_doc.add_paragraph(tempFile + '\n')#.add_run()



