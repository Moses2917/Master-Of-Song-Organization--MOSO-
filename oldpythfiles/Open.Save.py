import docx
from glob import glob
from docx.shared import Pt
my_doc = docx.Document("C:/Users/moses/OneDrive/Choir Songs Template.docx")

def process(filename):
    text = ""
    doc = docx.Document(filename)
    
    # # unzip the docx in memory
    # zipf = zipfile.ZipFile(docx)
    # filelist = zipf.namelist()

    for p in doc.paragraphs:
        # if doc.text #if doc.text has a tab(\t) then add a tab stop
        # p.paragraph_format
        # font = p.font
        tab_stops = p.paragraph_format.tab_stops
        if p is not None:
            text += '\t'
        text += p.text + "\n"

    # # get main text
    # doc_xml = 'word/document.xml'
    # text += xml2text(zipf.read(doc_xml))
    return text


filename = "C:/Users/moses/OneDrive/Երգարան Word Files/10000.docx"#gets name
doc = docx.Document(filename)
# style = doc.styles
tempFile = process(filename)
my_doc.add_paragraph(tempFile + '\n')
font = my_doc.styles['Normal'].font
font.name = 'Arial'
font.size = Pt(22)
my_doc.save("C:/Users/moses/OneDrive/Երգեր/11.2022/test.docx")
print("HI")

