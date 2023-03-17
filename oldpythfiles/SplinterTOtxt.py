############ONLY NESSESARY FOR 661 =>
from splinter import Browser
import docx
from docx.shared import Pt
failedSongs = []
# site_url = r"http://www.ergaran.in/2018/07/" + str(672) + ".html" 

def basicBR():
    # Open a web browser
    browser = Browser('firefox')

    # Visit the website
    browser.visit('http://www.ergaran.in/2018/07/662.html')

    # Find and extract the data
    data = browser.find_by_css(".post-body").text

    # Print the data
    print("662\n" + data)
    
    # Visit the website
    browser.visit('http://www.ergaran.in/2018/07/663.html')

    # Find and extract the data
    data = browser.find_by_css(".post-body").text

    # Print the data
    print("663\n" + data)
    # Close the web browser
    browser.quit()

basicBR()


songNum = 1
browser = Browser('firefox')

while songNum < 661:
    try:
        # Visit the website
        browser.visit('http://www.ergaran.in/2016/01/' + str(songNum) + '.html')

        # Find and extract the data
        data = browser.find_by_css(".post-body").text

        # Print the data
        print(str(songNum)+ "\n" + data)
        
        #Store data in 
        text = str(songNum) + "\n"
        text += data
        
        # f.write(text)
        doc = docx.Document()
        doc.add_paragraph(text)
        font = doc.styles['Normal'].font
        font.name = 'Arial'
        font.size = Pt(22)
        doc.save("D:/webSong/" + str(songNum) + ".docx")
        # f.close
    except:
        print("Did not find: " + str(songNum))
        failedSongs.append(songNum)
    songNum += 1

def theRest():
    songNum=1
    while songNum < 941:
        try:
            # Visit the website
            browser.visit('http://www.ergaran.in/2018/07/' + str(songNum) + '.html')

            # Find and extract the data
            data = browser.find_by_css(".post-body").text

            # Print the data
            print(str(songNum)+ "\n" + data)
            
            #Store data in 
            text = str(songNum) + "\n"
            text += data
            
            # f.write(text)
            doc = docx.Document()
            doc.add_paragraph(text)
            font = doc.styles['Normal'].font
            font.name = 'Arial'
            font.size = Pt(22)
            doc.save("D:/webSong/" + str(songNum) + ".docx")
            # f.close
        except:
            print("Did not find: " + str(songNum))
            failedSongs.append(songNum)
        songNum += 1

    while songNum < 1000:
        try:
            # Visit the website
            browser.visit('http://www.ergaran.in/2018/08/' + str(songNum) + '.html')

            # Find and extract the data
            data = browser.find_by_css(".post-body").text

            # Print the data
            print(str(songNum) + "\n" + data)

            # Store data in
            text = str(songNum) + "\n"
            text += data
            doc = docx.Document()
            doc.add_paragraph(text)
            font = doc.styles['Normal'].font
            font.name = 'Arial'
            font.size = Pt(22)
            doc.save("D:/webSong/" + str(songNum) + ".docx")
            # f.close
        except:
            print("Did not find: " + str(songNum))
            failedSongs.append(songNum)
        songNum += 1

browser.quit()

f = open("D:/webSong/" + "failedSongs" + ".txt", 'w', encoding='utf_8')
for x in failedSongs:
    print(x)
    if failedSongs.index(x) < len(failedSongs)-1:
        f.write(str(x)+",")
    else:
        f.write(str(x))
f.close
