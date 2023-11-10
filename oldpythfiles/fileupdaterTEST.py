import re, docx
from docx.shared import Pt
def myWay():
    import re, time, docx, json
    from os import path as pth, remove
    from docx.shared import Pt

    def getDocText(filename):
        text = ""

        doc = docx.Document(filename)
        for p in doc.paragraphs:
            text += p.text + "¶\n"
        return text

    song_doc = getDocText("D:/04.06.23.docx")
    print(song_doc)
    doc = docx.Document()
    doc.add_paragraph(song_doc)
    doc.save("D:/savedFile.docx")
    
# def theEdgeWay():
# import docx
# from docx.shared import Pt

# def getDocTextAndIndentation(filename):
#     doc = docx.Document(filename)
#     text_and_indentation = []
#     for p in doc.paragraphs:
#         first_line_indent = p.paragraph_format.first_line_indent
#         left_indent = p.paragraph_format.left_indent
#         right_indent = p.paragraph_format.right_indent
#         text_and_indentation.append({
#             'text': p.text,
#             'first_line_indent': first_line_indent,
#             'left_indent': left_indent,
#             'right_indent': right_indent
#         })
#     return text_and_indentation
def getDocText(filename):
    text = ""

    doc = docx.Document(filename)
    for p in doc.paragraphs:
        text += p.text + "\n"
    return text

# def getDocTextAndIndentation(filename, pattern):
#     matches = re.findall("\[start:song](.*?)\[end:song]",song_doc,re.DOTALL)

#     for match in matches:
#         print(match)

#     oldMatches = re.findall("\[start:song:old](.*?)\[end:song:old]",song_doc,re.DOTALL)

#     for oldMatch in oldMatches:
#         print(oldMatch)
#     doc = docx.Document(filename)
#     text_and_indentation = []
#     matches = re.findall(pattern, (doc.paragraphs).text,re.DOTALL)
#     for p in doc.paragraphs:
#         #these next three lines save the places where there is an indent
#         first_line_indent = p.paragraph_format.first_line_indent
#         left_indent = p.paragraph_format.left_indent
#         right_indent = p.paragraph_format.right_indent
#         # matches = re.findall(pattern, p.text,re.DOTALL)
#         for match in matches:
#             text_and_indentation.append({
#                 'text': match,
#                 'first_line_indent': first_line_indent,
#                 'left_indent': left_indent,
#                 'right_indent': right_indent
#             })
#     return text_and_indentation
def getDocTextAndIndentation(filename):
    doc = docx.Document(filename)
    text_and_indentation = [] #turn into a list of lists
    song = []
    bookOld = False
    for p in doc.paragraphs:
        if "[start:song" in p.text:
            song = []
            if "old" in p.text:
                bookOld = True
        if not("end" in p.text or "start" in p.text):
            first_line_indent = p.paragraph_format.first_line_indent
            left_indent = p.paragraph_format.left_indent
            right_indent = p.paragraph_format.right_indent
            song.append({
                'text': p.text,
                # 'book': re.findall(pattern, p.text,re.DOTALL)[0],
                'old': bookOld,
                'first_line_indent': first_line_indent,
                'left_indent': left_indent,
                'right_indent': right_indent
            })
        if "end" in p.text:
            #push song to text var and reset song var
            text_and_indentation.append(song)
            bookOld = False
            song = []
    return text_and_indentation

def FindName(song_text):
    return (re.findall("\d+",song_text[0]['text']))[0] + str(song_text[0]['old'])

def createDocFromTextAndIndentation(text_and_indentation):
    doc = docx.Document()
    for song in text_and_indentation:
        for paragraph_info in song:
            p = doc.add_paragraph(paragraph_info['text'])
            print(paragraph_info['text'])
            p.paragraph_format.space_after = 0
            if paragraph_info['first_line_indent'] is not None:
                p.paragraph_format.first_line_indent = paragraph_info['first_line_indent']
            if paragraph_info['left_indent'] is not None:
                p.paragraph_format.left_indent = paragraph_info['left_indent']
            if paragraph_info['right_indent'] is not None:
                p.paragraph_format.right_indent = paragraph_info['right_indent']
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(22)
        s = FindName(song)
        doc.save("D:/"+s+".docx")
    
input_filename = 'D:/04.06.23.docx'
text_and_indentation = getDocTextAndIndentation(input_filename)
createDocFromTextAndIndentation(text_and_indentation)