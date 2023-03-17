from html import parser
import time, os.path, bs4, requests
from xml.etree.ElementInclude import include
from time import sleep
from bs4 import BeautifulSoup
import docx
from docx.shared import Pt
bs4.BeautifulSoup()
songNum = 770
webpage = r"http://www.ergaran.in/2018/07/" + str(songNum) + ".html"
reqs = requests.get(webpage)
borsht = BeautifulSoup(reqs.content.decode('utf-8'), 'html.parser')
# print(borsht.prettify())
# print(borsht.find(class_ = 'post hentry uncustomized-post-template').get_text(" \n ", strip=True)) ## it works!!!
# print(borsht.find(class_ = 'post-body entry-content').get_text("\n")) ## it works!!!
Newborsht = BeautifulSoup(reqs.content.decode('utf-8'), 'html5lib')
print("\n")
# Newborsht.prettify()
bigBoySTR = Newborsht.find(class_ = 'post-body entry-content').get_text("\n")
print(Newborsht.find(class_ = 'post-body entry-content').get_text("\n")) #Useless

failedSongs = []
songNum = 661
# docs = docx.Document("D:/Ergaran Web Word Songs/BaseDoc.docx")

while songNum < 661:
    try:
        webpage = r"http://www.ergaran.in/2016/01/" + str(songNum) + ".html"
        reqs = requests.get(webpage)
        borsht = BeautifulSoup(reqs.text, "lxml")
        # f = open("D:/Ergaran Web Word Songs/"+ str(songNum) + ".txt", 'w', encoding='utf_8')
        print(borsht.find(class_ = "post-body entry-content").get_text("")) ## it works!!!
        text = str(songNum)
        text += borsht.find(class_ = "post-body entry-content").get_text("")
        # f.write(text)
        doc = docx.Document()
        doc.add_paragraph(text)
        font = doc.styles['Normal'].font
        font.name = 'Arial'
        font.size = Pt(22)
        doc.save("D:/webSong/"+ str(songNum) + ".docx")
        # f.close
    except:
        print("Did not find: " + str(songNum))
        failedSongs.append(songNum)
    songNum += 1


while songNum < 941:
    try:
        webpage = r"http://www.ergaran.in/2018/07/" + str(songNum) + ".html"
        reqs = requests.get(webpage)
        borsht = BeautifulSoup(reqs.text, "lxml")
        # f = open("D:/Ergaran Web Word Songs/"+ str(songNum) + ".txt", 'w', encoding='utf_8')
        print(str(songNum) + borsht.find(class_ = "post-body entry-content").get_text("")) ## it works!!!
        text = str(songNum)
        text += borsht.find(class_ = "post-body entry-content").get_text("")
        # f.write(text)
        doc = docx.Document()
        doc.add_paragraph(text)
        font = doc.styles['Normal'].font
        font.name = 'Arial'
        font.size = Pt(22)
        doc.save("D:/webSong"+ str(songNum) + ".docx")
        # f.close
    except:
        print("Did not find: " + str(songNum))
        failedSongs.append(songNum)
    songNum += 1

while songNum < 1000:
    try:
        webpage = r"http://www.ergaran.in/2018/08/" + str(songNum) + ".html"
        reqs = requests.get(webpage)
        borsht = BeautifulSoup(reqs.text, "lxml")
        # f = open("D:/Ergaran Web Word Songs/"+ str(songNum) + ".txt", 'w', encoding='utf_8')
        print(borsht.find(class_ = "post-body entry-content").get_text("\n")) ## it works!!!
        text = str(songNum)
        text += borsht.find(class_ = "post-body entry-content").get_text("\n")
        # f.write(text)
        doc = docx.Document()
        doc.add_paragraph(text)
        font = doc.styles['Normal'].font
        font.name = 'Arial'
        font.size = Pt(22)
        doc.save("D:/webSong/"+ str(songNum) + ".docx")
        # f.close
    except:
        print("Did not find: " + str(songNum))
        failedSongs.append(songNum)
    songNum += 1


f = open("D:/webSong/"+ "failedSongs" + ".txt", 'w', encoding='utf_8')
for x in failedSongs:
    print(x)
    if failedSongs.index(x) < len(failedSongs)-1:
        f.write(str(x)+",")
    else:
        f.write(str(x))
f.close

## So far failedsongs are [62,94,95,96,704,712,734,751,909]

##    sleep(0.1)


# failedSongs = []

# while songNum < 659: ## 2016/01/
#     webpage = r"http://www.ergaran.in/2016/01/" + str(songNum) + ".html"
#     try:
#         reqs = requests.get(webpage)
#         sleep(0.5)
#         borsht = BeautifulSoup(reqs.content, "html.parser")
#         text = borsht.find(class_ = 'post hentry uncustomized-post-template').text
#         print(text) ## it works!!!
#         f = open("M:/newWord/"+ str(songNum) + ".txt", 'w', encoding='utf_8')
#         sleep(0.5)
#         f.write(text)
#         f.close
#     except:
#         failedSongs.append(songNum)
#     songNum += 1

# while songNum < 941: ## 2018/07/
#     webpage = r"http://www.ergaran.in/2018/07/" + str(songNum) + ".html"
#     try:
#         reqs = requests.get(webpage)
#         sleep(0.5)
#         borsht = BeautifulSoup(reqs.content, "html.parser")
#         text = borsht.find(class_ = 'post hentry uncustomized-post-template').text
#         print(text) ## it works!!!
#         f = open("M:/newWord/"+ str(songNum) + ".txt", 'w', encoding='utf_8')
#         sleep(0.5)
#         f.write(text)
#         f.close
#     except:
#         failedSongs.append(songNum)
#     songNum += 1

# while songNum < 1001: ## 2018/08/
#     webpage = r"http://www.ergaran.in/2018/08/" + str(songNum) + ".html"
#     try:
#         reqs = requests.get(webpage)
#         sleep(0.5)
#         borsht = BeautifulSoup(reqs.content, "html.parser")
#         text = borsht.find(class_ = 'post hentry uncustomized-post-template').text
#         print(text) ## it works!!!
#         f = open("M:/newWord/"+ str(songNum) + ".txt", 'w', encoding='utf_8')
#         sleep(0.5)
#         f.write(text)
#         f.close
#     except:
#         failedSongs.append(songNum)
#     songNum += 1
