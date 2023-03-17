import time, docx, os, re
from docx.shared import Pt
from glob import glob
## Stuff for doc2txty
import xml.etree.ElementTree as ET
## use a try method to iterate through already made word files, and just in case go to the red words folder
my_doc = docx.Document("C:/Users/moses/OneDrive/Choir Songs Template.docx")

time.strftime('%B')

wordFile = open("Choir Songs Template.dotx", 'r', encoding='utf_8')
template = wordFile.read()
wordFile.close
month = time.strftime('%m')
year = time.strftime('%y')
fullYear = time.strftime('%Y')
day = time.strftime('%d')
# word = open("C:/Users/moses/OneDrive/Երգեր/" + month + "." + fullYear + "/" + month + "." + day + ".P" + year + ".docx", 'a', encoding='utf_8')
# word.write(template)

def ask():
    print("")

songs = [] # For the song numbers
imp = [] # Stores important info about the location of the song, such as old(ppt/from other church) or new(red book)

book = input("Are all the songs from new/red songs\n y|n\n")
if("y" == book.lower()):
    while True:##Assumes that all songs are from new.
        imp.append("n")
        s = input("song num: ")
        if s == "":
            break
        else:
            songs.append(s)
else:
    print("Press any key except for 'n' or 'o' to cancel the loop")
    while True:
        inp = input("Enter 'n' or 'o' ")
        if inp.lower() == 'o' or inp.lower() == 'n':
            imp.append(inp)
            s = input("song num: ")
            songs.append(s)
            ##make a list/ array to store the values of which song num is old or new
        else:
            break
        
        


def process(filename):
    text = ""
    doc = docx.Document(filename)
    for p in doc.paragraphs:
        text += p.text + "\n"
    return text



for x in songs:
    y = songs.index(x)
    if 'n' in imp[y]:
        print(x + "\n")
        try:
            filename = glob("C:/Users/moses/OneDrive/Երգարան Word Files/" +str(x) + "*.docx")[0] #gets name
            tempFile = "[start:song]\n" + process(filename) + "[end:song]"
        except:
            tempFile = "[start:song]\n" + process("C:/Users/moses/OneDrive/RED Words/" + str(x) + ".docx") + "[end:song]"
    else:
        print(x + "\n")
        try:
            filename = glob("C:/Users/moses/OneDrive/Word songs/" + x + "*")[0] # gets name
            doc = docx.Document(filename)
            tempFile = "[start:song:old]\n" + process(filename) + "[end:song:old]" #9168588420 #
        except:
            print(x + ":" + " FileNotFoundError")
            tempFile = ""
        
    my_doc.add_paragraph(tempFile + '\n')#.add_run()



print("Adjusting for readability....")
font = my_doc.styles['Normal'].font
font.name = 'Arial'
font.size = Pt(22)
print("it is done")
# my_doc.save("C:/Users/moses/OneDrive/Երգեր/" + month + "." + fullYear + "/" + month + "." + day + "." + year + "TEST FILE DO NOT CLICK" + ".docx")
my_doc.save("C:/Users/moses/OneDrive/Երգեր/" + month + "." + fullYear + "/" + month + "." + day + "." + year + ".docx")






# # songs: 350 251
# # imp:    o  n
# for x in songs:
#     y = songs.index(x)
#     if 'n' in imp[y]:
#         print(x + "\n")
#         try:
#             filename = glob("C:/Users/moses/OneDrive/Երգարան Word Files/" + x + "*.docx")[0] #gets name
#             doc = docx.Document(filename) ##Add a way to double check that the song is the right one
#             # style = doc.styles            and not 691 pretending to be 69 ok ok
#             tempFile = process(filename)
#         except:
#             songFile = open("C:/Users/moses/OneDrive/Ergaran Web Word Songs/" + x + ".txt", 'r', encoding='utf_8')
#             tempFile = songFile.read()
#             songFile.close()
#     else:
#         print(x + "\n")
#         try:
#             filename = glob("C:/Users/Armne/OneDrive/Word songs/" + x + "*")[0] # gets name 
#             doc = docx.Document(filename)
#             tempFile = process(filename)
#         except:
#             print(x + ":" + " FileNotFoundError")
#             tempFile = ""
        
#     # word.write("\n" + tempFile + '\n')
#     my_doc.add_paragraph(tempFile + '\n')#.add_run()
