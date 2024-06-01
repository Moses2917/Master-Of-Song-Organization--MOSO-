import chromadb
from os import environ as env
from docx import Document
from chromadb.utils import embedding_functions
# sentance_transformer_ef = embedding_functions.SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L12-v2")
client = chromadb.PersistentClient(path="{}\\documents\\code\\python\\.chroma".format(env.get("OneDrive")))
# try:
#     pass
#     collection = client.create_collection(name="songs")
# except:
#     collection = client.get_collection(name="songs")
# collection = client.get_collection(name="songs")
collection = client.get_or_create_collection(name="indivPara")#, embedding_function=sentance_transformer_ef)
from json import load
with open('wordSongsIndex.json','r',encoding='utf-8') as f: #Do this one, and also "REDergaran"
    song_index = load(f)
    # book = 'Old'
didnt_find = []
for songNum in song_index['SongNum']:
    try:
        filePath = env.get("OneDrive")+"\\"+song_index['SongNum'][songNum]['latestVersion']
        doc = Document(filePath)
        word_doc = ""
        # from hashlib import sha256
        # from random import random
        # for p in doc.paragraphs:
        #     word_doc = p.text + songNum + "Old" + str(round(random()*100000))
        #     id = sha256(word_doc.encode()).hexdigest()
        #     collection.add(
        #         documents=p.text,
        #         metadatas=[{
        #             'book': "Old",
        #             'songnum': songNum,
        #         }],
        #         ids=id
        #     )
        from random import random
        for p in doc.paragraphs:
            word_doc = word_doc + p.text
        word_doc = word_doc
        from hashlib import sha256
        id = sha256((
                            word_doc + songNum + "Old"+ str(round(random()*100000))

                    ).encode()).hexdigest()
        collection.add(
            documents=word_doc,
            metadatas=[{
                'book': "Old",
                'songnum': songNum,
            }],
            ids=id
        )
    except:
        didnt_find.append(env.get("OneDrive")+"\\"+song_index['SongNum'][songNum]['latestVersion'])

with open("notFound.txt",'a',encoding='utf-8') as f:
    f.write(str(didnt_find))
print(collection.peek())

