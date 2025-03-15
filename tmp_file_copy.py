import os
from os.path import join

# doc_dir:str = "/Users/movsesmovsesyan/Documents/test_doc.docx"
doc_dir:str = r"C:\Users\moses\OneDrive\Երգեր\03.2025\03.16.25.docx"

def tmp_file_copy(doc_dir):
    import shutil
    import tempfile
    with tempfile.TemporaryDirectory() as tmp_dir:
        filename = doc_dir.split('/')[-1] # gets the filename
        tmp_file = os.path.join(filename, tempfile.tempdir)
        temp_file_path = shutil.copy(doc_dir, tmp_file)
        # Perform rest of ops in the 'with' statement

# tmp_file_copy(doc_dir)

import concurrent
from docx import Document

def tmp_file_copy(doc_dir):
    import tempfile
    import subprocess
    with tempfile.TemporaryDirectory() as tmp_dir:
        filename = os.path.basename(doc_dir) # gets the filename
        tmp_file = join(tmp_dir, filename)
        temp_file_path: str = subprocess.run(['copy', f"{doc_dir}", f"{tmp_dir}"], shell=True,)
        # Perform rest of ops in the 'with' statement
        return Document(tmp_file)

def async_call():
    import docx
    doc = tmp_file_copy(doc_dir) #Document(filePth)  # load doc file
    docParagraphs = doc.paragraphs
    print(docParagraphs)

from concurrent.futures import ThreadPoolExecutor
with ThreadPoolExecutor() as futures:
    futures.submit(async_call).result()
    