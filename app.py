from datetime import datetime
from glob import glob
from concurrent.futures import ThreadPoolExecutor, thread
from os import environ as env
from os import stat
import os
from os.path import join
from pprint import pprint
from random import random
import re
import secrets
from urllib.parse import quote_plus, urlencode
from authlib.integrations.flask_client import OAuth
from dotenv import find_dotenv, load_dotenv
from flask import Flask, jsonify, render_template, request, redirect, url_for, session, flash
import json
from concurrent.futures import ThreadPoolExecutor
#Import Custom Lyrics Search Engine
from lyric_search_engine import SearchEngine
import logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
ENV_FILE = find_dotenv("{}\Documents\Code\.env".format(env.get("OneDrive")))
if ENV_FILE:
    load_dotenv(ENV_FILE)

app = Flask(__name__)

secret_key = env.get("APP_SECRET_KEY")

app.secret_key = secret_key if secret_key else secrets.token_urlsafe(16)

oauth = OAuth(app)

oauth.register(
    "auth0",
    client_id=env.get("AUTH0_CLIENT_ID"),
    client_secret=env.get("AUTH0_CLIENT_SECRET"),
    client_kwargs={
        "scope": "openid profile email",
    },
    server_metadata_url=f'https://{env.get("AUTH0_DOMAIN")}/.well-known/openid-configuration'
)

search_engine = SearchEngine()
song_lyrics = search_engine.load_json_data('AllLyrics.json')

onedrive_path = env.get("OneDrive") if os.name != "posix" else "/Users/movsesmovsesyan/Library/CloudStorage/OneDrive-Personal"

def open_past_songs():
    with open("songs_cleaned.json" , 'r', encoding='utf-8') as f:
        all_past_songs:dict = json.load(f)
    return all_past_songs

all_past_songs = open_past_songs()
# Should only load this once instead of every time
with open('REDergaran.json', mode='r', encoding='utf-8') as json_file:
    REDergaran:dict = json.load(json_file)
with open('wordSongsIndex.json', mode='r', encoding='utf-8') as json_file:
    wordSongsIndex:dict = json.load(json_file)

@app.route("/robots.txt", methods=["GET"])
def robots():
    with open('./templates/robots.txt', 'r') as f:
        return f.read() 

#This route is responsible for actually saving the session for the user, so when they visit again later, they won't have to sign back in all over again.
@app.route("/callback", methods=["GET", "POST"])
def callback():
    """
    Handles the callback from the authorization server.

    This function is responsible for exchanging the authorization code for an access token,
    storing the user's session, and redirecting the user to the temporary home page.

    Parameters:
    None

    Returns:
    redirect: A redirect response to the temporary home page.
    """ 
    token = oauth.auth0.authorize_access_token()
    session["user"] = token
    session['user']['userinfo']['admin'] = isUserAllowed(session['user']['userinfo']['email'])
    # print(session["user"])
    #remember to flash('Բարի Գալուստ, {{session.user.userinfo.name}}!')
    # return redirect("/")
    return redirect("/")
  

#the /login route, users will be redirected to Auth0 to begin the authentication flow.
@app.route("/login")
def login():
    """
    This function defines the /login route, which redirects users to Auth0 to begin the authentication flow.
    It does not take any parameters and returns the authorization redirect response from Auth0.
    """
    return oauth.auth0.authorize_redirect(
        redirect_uri=url_for("callback", _external=True)
    )

#this route handles signing a user out from your application.
@app.route("/logout")
def logout():
    """
    Handles user logout by clearing the session and redirecting to Auth0's logout endpoint.

    Returns:
        redirect: A redirect response to Auth0's logout endpoint.
    """
    session.clear()
    return redirect(
        "https://" + env.get("AUTH0_DOMAIN")
        + "/v2/logout?"
        + urlencode(
            {
                "returnTo": url_for("home", _external=True),
                "client_id": env.get("AUTH0_CLIENT_ID"),
            },
            quote_via=quote_plus,
        )
    )

# Function to check the eligability of user to have admin access
def isUserAllowed(email):
    """
    Checks if the provided email is allowed to access the webpage.
    
    Args:
        email (str): The email to be checked.
    
    Returns:
        bool: True if the email is allowed, False otherwise.
    """
    with open("{}\\Documents\\Code\\allowedEmails.csv".format(env.get("OneDrive")), 'r') as f:
        return email in f.read()

# Table data loading logic
def load_table_data(book:str):
    """Returns a whole json dict of either old or new, based upon input\n
    Reads the respective index file for the book given, and will return a dict

    Args:
        book (str): name of book, can be REDergaran or wordSongsIndex

    Returns:
        dict: a dict containing all of the songs in that index
    """
    try:
        return wordSongsIndex.get('SongNum') if book == "wordSongsIndex" else REDergaran.get('SongNum')
    except FileNotFoundError:
        return None

# @app.route('/getTableData', methods=['GET', 'POST'])
def getSong(book:str, songnum:str, batch = 0) -> dict:
    """
    Retrieves a song from a JSON file based on the provided book and song number. The options are 'old', 'new', 'redergaran', and 'wordsongsindex'.

    Args:
        book (str): The book from which to retrieve the song.
        songnum (str): The number of the song to retrieve.
        batch (int, optional): The batch number. Defaults to 0.

    Returns:
        dict: A dictionary containing the song data.
    """
    if batch == 0:
        from json import load
        if book.lower() == "old" or book.lower() == "wordsongsindex":
            # with open("wordSongsIndex.json", 'r', encoding='utf-8') as f:
            #     wordSongs = load(f)["SongNum"]
            #     return wordSongs[songnum]
            return wordSongsIndex["SongNum"][songnum]
        else:
            # with open("REDergaran.json", 'r', encoding='utf-8') as f:
            #     REDergaran = load(f)["SongNum"]
            #     return REDergaran[songnum]
            return REDergaran["SongNum"][songnum]
    else:
        pass

def openWord(songNum, book):
    """
    Opens a word document based on the provided song number and book, 
    and returns the contents of the document as a HTML formatted string.

    Args:
        songNum (int): The number of the song to be opened.
        book (str): The name of the book containing the song.

    Returns:
        str: A HTML formatted string containing the contents of the word document.
    """
    from docx import Document
    #Used to find and open word doc, sends back a html formated str of it
    
    if ("word" in book.lower() or "old" in book.lower()):
        index = wordSongsIndex
    else:
        index = REDergaran
    
    if index["SongNum"].get(songNum,None):
        songPth = index["SongNum"][songNum]["latestVersion"]
        filePth = env.get("OneDrive")+"\\"+songPth #find pth from index, and attach the location for onedrive
        doc = Document(filePth) #load doc file
        docParagraphs = doc.paragraphs # returns a list of doc paragrpahs from which text will be extracted
        text = ''
        
        for para in docParagraphs:
            text += para.text + '\n'
        
        # Split the text into chunks based on line breaks
        chunks = text.split('\n\n')
        
        # Convert chunks to HTML
        html_chunks = []
        for chunk in chunks:
            lines = chunk.split('\n')
            html_lines = ['<p>' + line + '</p>' for line in lines]
            html_chunk = ''.join(html_lines)
            html_chunks.append(html_chunk)

        # Join the chunks with line breaks, adding or subtracting br will add or subtract the breaks between the paragraphs
        html_text = '<br>'.join(html_chunks)
        return html_text

def saveHtml(filePth, WordDoc):
    """
    Saves the contents of a Word document to an HTML file to later be displayed.

    Parameters:
        filePth (str): The path to the Word document file.
        WordDoc (str): The name of the Word document, not the path, but just the filename.

    Returns:
        None
    """
    from docx import Document
    from docx import document
    def tmp_file_copy(doc_dir) -> document.Document:
        import tempfile
        import shutil
        with tempfile.TemporaryDirectory() as tmp_dir:
            filename = os.path.basename(doc_dir) # gets the filename
            tmp_file = join(tmp_dir, filename)
            # if os.name == 'nt':
            #     temp_file_path = shutil.copyfile(doc_dir, tmp_dir)
            temp_file_path = shutil.copy(doc_dir, tmp_dir)
            # Perform rest of ops in the 'with' statement
            return Document(tmp_file)
    
    # print("Saving File...")
    
    doc = tmp_file_copy(filePth)
    from doc_color import get_colored_text, get_all_colored_text

    # Extract text WITH color info using Method 1's song logic
    text_with_colors = ''
    start_song = False
    song_counter = 0
    song_nums = []

    for para in doc.paragraphs:
        para_html = ''
        for run in para.runs:
            text = run.text

            if start_song:
                if song_counter != 1:
                    para_html += f'</div>'
                para_html += f'<div id="song-{song_counter}">'
                song_nums.append(text)
                start_song = False
                
            if "start" in text:
                start_song = True
                song_counter += 1
            
            # Add color formatting
            text_color = run.font.color.rgb
            if bool(text_color) and str(text_color) != "000000":
                para_html += f'<span style="color: #{text_color};">{text}</span>'
            else:
                para_html += text
        
        text_with_colors += para_html + '\n'

    # Spacing logic
    chunks = text_with_colors.split('\n\n')
    html_chunks = []

    for chunk in chunks:
        lines = chunk.split('\n')
        html_lines = []
        
        for line in lines:
            html_lines.append('<p>' + line + '</p>')
        
        html_chunk = ''.join(html_lines)
        html_chunks.append(html_chunk)

    html_text = '<br>'.join(html_chunks)

    with open(f"{onedrive_path}/Documents/Code/Python/htmlsongs/{WordDoc}.txt", 'w', encoding='utf-8') as f:
        f.write(html_text)
    # print(song_nums)
    return song_nums

@app.route('/search/<searchLyrics>', methods=['GET'])
def songSearch(searchLyrics) -> list:
    """
    Searches for songs based on provided lyrics and returns a list of search results.

    Args:
        searchLyrics (str): The lyrics to search for in the songs.

    Returns:
        list: A list of search results, where each result is a dictionary containing the book, song number, and title of the song.
    """

    if searchLyrics:
        data = request.get_json(silent=True, cache=True)
        num_results = 10 if not data.get("number_results", None) else data.get("number_results", None)
        # print(f"num Results = {num_results}")
        searchResults = []
        results = search_engine.search(query=searchLyrics, top_k=num_results) if num_results else search_engine.search(query=searchLyrics)
 
        # print(results)
        if results[0][2] > 0:
            for result in results:
                if result[0].lower() == 'old':
                    with open('wordSongsIndex.json', 'r', encoding='utf-8') as f:
                        index = json.load(f)
                else:
                    with open('REDergaran.json', 'r', encoding='utf-8') as f:
                        index = json.load(f) 
                
                title = index['SongNum'][result[1]]['Title']
                
                title = title.split('\n')[0]
                
                
                
                searchResults.append(
                    f'''<a class="list-group-item list-group-item-action" href="{url_for('display_song',book=result[0],songnum=result[1])}">{result[1]}: {title}</a>'''
                )
        # for result in searchResults:
        #     re.match
        return json.dumps(searchResults)

@app.route('/song/<book>/<songnum>', methods=['GET','POST'])
def display_song(book, songnum) -> str:
    """
    This function handles HTTP requests to the '/song/<book>/<songnum>' route.
    
    It takes two parameters: 'book' and 'songnum', which are used to identify a specific song.
    
    The function returns a rendered HTML template ('song.html') with the song's lyrics, 
    past songs, and similar songs.
    
    Parameters:
    book (str): The book identifier for the song (e.g., 'Old' or 'New').
    songnum (str): The song number identifier.
    
    Returns:
    A rendered HTML template with the song's lyrics, past songs, and similar songs.
    """
    from scanningDir import songSearch
    with open('wordSongsIndex.json', 'r', encoding='utf-8') as f:
        wordSongsIndex = json.load(f)

    with open('REDergaran.json', 'r', encoding='utf-8') as f:
        REDergaran = json.load(f)
    
    with open('song_occurrences.json', 'r', encoding='utf-8') as f:
        occr = json.load(f)
    
    book = book.lower()
    if book == 'wordsongsindex' or book == 'old':
        book = 'Old'
    else:
        book = 'New'
    past_songs = songSearch(songnum, book)
    similar_songs = None
    if songnum in occr[book]:
        similar_songs = occr[book][songnum]
    
    # template = f'''<a class="list-group-item list-group-item-action" href="{url_for('display_song',book=book,songnum=songnum)}">{songnum}: {title}</a>'''
    if past_songs != None:
        for songs in past_songs:
            song_titles = []
            for song_pair in songs['songs']:
                if not (None in song_pair):
                    # each song is a tuple ie: ('Old', '495')
                    # here I am simply adding a tuple(list(title))
                    title = ""
                    if song_pair[0] == 'Old':
                        if song_pair[1] in wordSongsIndex['SongNum']:
                            title = wordSongsIndex['SongNum'][song_pair[1]]["Title"]
                            title = title.split('\n')[0]
                    else:
                        if song_pair[1] in REDergaran['SongNum']:
                            title = REDergaran['SongNum'][song_pair[1]]["Title"]#REDergaran.get(['SongNum'][song_pair[1]]["Title"],None) # doing this to try to account for unusual song nums such as '32121'
                            title = title.split('\n')[0]
                    song_titles.append(f'''<a class="list-group-item list-group-item-action" href="{url_for('display_song',book=song_pair[0],songnum=song_pair[1])}">{song_pair[1]}: {title}</a>''')
            songs['songs'] = song_titles
    if similar_songs != None:
        song_titles = []
        for song_pair in similar_songs:
            # print(song_pair)
            song_pair = tuple(eval(song_pair))
            # print(song_pair)
            if not (None in song_pair):
                title = ""
                if song_pair[0] == 'Old':
                    if song_pair[1] in wordSongsIndex['SongNum']:
                        title = wordSongsIndex['SongNum'][song_pair[1]]["Title"]
                        title = title.split('\n')[0]
                else:
                    if song_pair[1] in REDergaran['SongNum']:
                        title = REDergaran['SongNum'][song_pair[1]]["Title"]#REDergaran.get(['SongNum'][song_pair[1]]["Title"],None) # doing this to try to account for unusual song nums such as '32121'
                        title = title.split('\n')[0]
                song_titles.append(f'''<a class="list-group-item list-group-item-action" href="{url_for('display_song',book=song_pair[0],songnum=song_pair[1])}">{song_pair[1]}: {title}</a>''')
        similar_songs = song_titles
        # print(similar_songs)
    
    import concurrent.futures
    with concurrent.futures.ThreadPoolExecutor() as exec:
        future = exec.submit(openWord,songnum,book)
        lyrics = future.result()
        
    return render_template('song.html', lyrics=lyrics, past_songs=past_songs, similar_songs=similar_songs, songnum=songnum)

def save_json(json:dict, path:str):
    from json import dump
    with open(path, 'w', encoding='utf-8') as f:
        dump(json, f, ensure_ascii=False, indent=4)

@app.route('/today', methods=['GET'])
def today_songs():
    all_past_songs = open_past_songs() # Need to always get a fresh version of this
    latest_song = list(all_past_songs.items())[-1]
    song_dict:dict = latest_song[1] # the info stored inside the dictionary. ie: "03.16.25.docx": { "dateMod": 1741926049.0, "path": ... , "basePth": "Երգեր\\03.2025", "songList": str(list(tuple())) }
    last_modified_date: float = song_dict["dateMod"]
    WordDoc = latest_song[0]
    cached_txt_files = glob("htmlsongs\\"+WordDoc+"*") # Cached txt files
    basePth:str = all_past_songs[WordDoc]['basePth']
    # songPth = songPth.split("OneDrive")[1] # bc of the way it's saved ie C:\Users\moses\OneDrive\Երգեր\06.2024\06.25.24.docx
    onedrive: str | None = env.get("OneDrive")
    songPth = os.path.join(onedrive,basePth, WordDoc)
    # from doc_color import get_colored_text
    # colored_text = get_colored_text(songPth)
    dateModOnFile: datetime = datetime.fromtimestamp(last_modified_date)
    currDateMod: datetime = datetime.fromtimestamp(stat(songPth).st_mtime)
    # if currDateMod <= dateModOnFile:
    if False:
        # Bigger number means further in time
        # Run this if the date on file, is the latest date
        if cached_txt_files:
            with open(cached_txt_files[0], 'r', encoding='utf-8') as f:
                lyrics = f.read()
            return render_template("display_docx.html", lyrics = lyrics, colored_text=colored_text)
        else:
            try:
                with ThreadPoolExecutor() as futures:
                    future = futures.submit(saveHtml, songPth, WordDoc)
                with open(f"htmlsongs\\{WordDoc}.txt", 'r', encoding='utf-8') as f:
                    html_text = f.read()
                return render_template("display_docx.html", lyrics=html_text, colored_text=colored_text)
            except:
                with open(cached_txt_files[0], 'r', encoding='utf-8') as f:
                    lyrics = f.read()
                return render_template("display_docx.html", lyrics = lyrics, colored_text=colored_text)
    else:
        # all_past_songs[WordDoc]["dateMod"] = currDateMod.timestamp()
        # print(songPth)
        with ThreadPoolExecutor() as futures:
            future = futures.submit(saveHtml, songPth, WordDoc)
            # save = futures.submit(save_json, all_past_songs, "songs_cleaned.json")
            result = future.result()
            # result2 = save.result()
        with open(f"htmlsongs\\{WordDoc}.txt", 'r', encoding='utf-8') as f:
            html_text = f.read()
        return render_template("display_docx.html", lyrics=html_text, song_nums=result)

@app.route('/events', methods=["GET", "POST"])
def event(filename = r"Երգեր/Պենտեկոստե/2025/Պենտեկոստե.docx"):
    folder_path = os.path.join(onedrive_path,"Երգեր/Պենտեկոստե")
    print(folder_path)
    # filename = os.path.join(onedrive_path, filename)
    # from doc_color import get_colored_text
    # colored_text = get_colored_text(filename)
    if request.method == 'GET':
        # with ThreadPoolExecutor() as futures:
        #     future = futures.submit(saveHtml, filename, "Պենտեկոստե.docx") 
        #     # save = futures.submit(save_json, all_past_songs, "songs_cleaned.json")
        #     result = future.result()
        #     # result2 = save.result()
        # with open(f"htmlsongs/Պենտեկոստե.docx.txt", 'r', encoding='utf-8') as f:
        #     html_text = f.read()
        # return render_template("display_docx.html", lyrics=html_text, song_nums=result)
        # list to hold all dirs, with relative reference starting at fp
        roots = []
        os_files = {}
        for root, dirs, files in os.walk(folder_path):
            # Ignore the folder itself. So, if "static" we don't want it!
            if root != folder_path:
                root = root.replace("\\","/")
                # print(root)
                roots.append(root)
                tmp_file = []
                files = list(map(lambda file: f'<button type="submit" name="selected_file" value="{os.path.join(root,file)}" class="btn btn-outline-light m-2">{file.split(".")[0]}</button>' if '.ppt' not in file else '', files))
                os_files[root] = files
        roots.sort(reverse=True)
        # print(roots)
        # filenames = list(map(lambda x: re.sub(r".docx",'',os.path.basename(x)), glob(fp+'*')))
        return render_template("event.html", os_files=os_files, roots=roots)
    else:
        selected_file = request.form.get(key='selected_file')
        # selected_file = 
        # is_dir = request.form.get('is_dir')
        # print(selected_file)
        if selected_file and os.path.exists(selected_file):
            basename = os.path.basename(selected_file) # basename returns the filename. Could've been done in saveHtml
            
            with ThreadPoolExecutor() as futures:
                future = futures.submit(saveHtml, selected_file, basename) 
                save = futures.submit(save_json, all_past_songs, "songs_cleaned.json")
                result = future.result()
                result2 = save.result()
            with open(f"htmlsongs\\{basename}.txt", 'r', encoding='utf-8') as f:
                html_text = f.read()
            return render_template("display_docx.html", lyrics=html_text)
        else:
            return 'No text found'


@app.route('/youth', methods=["GET", "POST"])
def youth():
    # folder_path = os.path.join(onedrive_path,"Youth Songs")

    
    folder_path = os.path.join(onedrive_path,"Youth Songs")
    # folder_path = "static/pictures"
    # filename = os.path.join(onedrive_path, filename)
    # from doc_color import get_colored_text
    # colored_text = get_colored_text(filename)
    if request.method == 'GET':
        roots = []
        os_files = {}
        for root, dirs, files in os.walk(folder_path):
            # Ignore the folder itself. So, if "static" we don't want it!
            # if root != folder_path:
            root = root.replace("\\","/")
            # print(root)
            roots.append(root)
            tmp_file = []
            files = list(map(lambda file: f'<button type="submit" name="selected_file" value="{os.path.join(root,file)}" class="btn btn-outline-light m-2">{file.split(".")[0]}</button>' if '.ppt' not in file else '', files))
            os_files[root] = files
        roots.sort(reverse=True)
        # print(roots)
        # filenames = list(map(lambda x: re.sub(r".docx",'',os.path.basename(x)), glob(fp+'*')))
        # pprint(os_files)
        return render_template("youth.html", os_files=os_files)#, roots=roots)
    else:
        selected_file = request.form.get(key='selected_file')
        # selected_file = 
        # is_dir = request.form.get('is_dir')
        # print(selected_file)
        if selected_file and os.path.exists(selected_file):
            basename = os.path.basename(selected_file) # basename returns the filename. Could've been done in saveHtml
            
            with ThreadPoolExecutor() as futures:
                future = futures.submit(saveHtml, selected_file, basename) 
                save = futures.submit(save_json, all_past_songs, "songs_cleaned.json")
                result = future.result()
                result2 = save.result()
            with open(f"htmlsongs\\{basename}.txt", 'r', encoding='utf-8') as f:
                html_text = f.read()
            return render_template("display_docx.html", lyrics=html_text)
        else:
            return 'No text found'


@app.route('/song/docx/<WordDoc>', methods=['GET','POST'])
def ServiceSongOpen(WordDoc) -> str:
    """
    Renders the song.html template with lyrics from a .docx file.

    Parameters:
    WordDoc (str): The name of the .docx file.

    Returns:
    str: The rendered song.html template with lyrics.

    Notes:
    - If the .docx file exists in the htmlsongs directory, it reads the lyrics from the file.
    - If the .docx file does not exist, it attempts to open the file from the OneDrive path and save the lyrics as an html file.
    - If the file does not exist and cannot be opened, it flashes an error message.
    """
    if '.docx' in WordDoc:
        foundFiles = glob("htmlsongs\\"+WordDoc+"*")
        if foundFiles:
            # print(foundFiles)
            with open(foundFiles[0], 'r', encoding='utf-8') as f:
                lyrics = f.read()
            return render_template("song.html", lyrics = lyrics)
        else:
            # with open("songs.json" , 'r', encoding='utf-8') as f:
            #     songs = json.load(f)
            
            songPth = all_past_songs[WordDoc]['path']
            songPth = songPth.split("OneDrive")[1] # bc of the way it's saved ie C:\Users\moses\OneDrive\Երգեր\06.2024\06.25.24.docx
            onedrive = env.get("OneDrive")
            songPth = onedrive+songPth
            # print(songPth)
            # songPth = fr"{onedrive}\Երգեր\Պենտեկոստե\2024\2024 Պենտեկոստե.docx"
            import threading as th
            wordDocThread = th.Thread(target=saveHtml,args=[songPth,WordDoc])#, args=[MS_WORD, songPth])
            wordDocThread.start()
            from time import sleep
            sleep(0.75)
            # saveHtml()
            with open(f"htmlsongs\\{WordDoc}.txt", 'r', encoding='utf-8') as f:
                html_text = f.read()
            return render_template("display_docx.html", lyrics=html_text)
    else:
        flash("That song does not exist",'error')

@app.route('/song/<book>/<songnum>/attributes', methods=['GET','POST'])
def getSongAttributes(book,songnum) -> dict:
    """
    Handles HTTP requests to retrieve song attributes.

    Parameters:
        book (str): The book from which the song is retrieved.
        songnum (str): The song number within the book.

    Returns:
        A JSON response containing the song attributes.
    """
    return jsonify(getSong(book, songnum))

@app.route('/attributeSearch', methods=['GET','POST'])
def attributeSearch() -> dict:
    """This inputs a list of attributes and the finds songs whose attributes match the input attributes.
    
    Input:
        dict: A list of attributes
        ex:
        songattrs = {
            "Comments": "Dance-P5-5",
            "Title": "Եղբայրնե՛ր, ցնծացե՛ք",
            "Worship_Song": "",
            "key": "Em",
            "latestVersion": "Երգարան Word Files/1 Եղբայրնե՛ր, ցնծացե՛ք.docx",
            "opening_Song": "",
            "song_type": "Opening Song",
            "speed": "105",
            "style": "Disco",
            "timeSig": "4/4",
            "v1": "Երգարան Word Files/1 Եղբայրնե՛ր, ցնծացե՛ք.docx"
        }
        attributes = {key: true, speed: true, style: false, song_type: false, timeSig: true}
        
        dict: A list of song attributes
    
    Returns:
        dict: A dictionary where the keys are the song numbers and the values are the dictionaries of the respective songs

    """
    data = request.get_json()
    attributes = data["attributes"]
    songattrs = data["songattrs"]
    # print("Songattrs:",songattrs)
    # print("attributes:",attributes)
    if songattrs.get("Comments", None):
        del songattrs["Comments"]
    if songattrs.get("Title", None):
        del songattrs["Title"]
    
    temp = {}
    for attribute in songattrs:
        if attributes.get(attribute):#Check if the attribute is in the dictionary
            if attributes[attribute]:
                temp[attribute] = songattrs[attribute]
    songattrs = temp
    # print(songattrs)
    from json import load
    # with open("wordSongsIndex.json", 'r', encoding='utf-8') as f:
    #     wordSongs = load(f)["SongNum"]
    
    # with open("REDergaran.json", 'r', encoding='utf-8') as f:
    #     REDergaran = load(f)["SongNum"]
    returnSongs = {}
    returnSongs["WordSongsIndex"] = {}
    returnSongs["REDergaran"] = {}
    def filter_songs(songs):
        found_songs = {}
        for songNum, song in songs.items():
            matched =True
            for attr in songattrs:
                foundSongAttr = song.get(attr)
                if foundSongAttr != songattrs[attr]: matched = False
            
            if matched:
                found_songs[songNum] = song

        return found_songs
                
    # Filter songs from both sources
    returnSongs["WordSongsIndex"] =filter_songs(wordSongsIndex["SongNum"])
    returnSongs["REDergaran"]=filter_songs(REDergaran["SongNum"])
    return jsonify(returnSongs)
            
def get_my_ip() -> str:
    """
    Retrieves the public IP address of the device making the request.
    
    Returns:
        str: The public IP address as a string.
    """
    from requests import get
    def get_local_ip():
        import socket
        try:
            # Create a socket object
            s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            # Doesn't need to be reachable
            s.connect(('8.8.8.8', 80))
            local_ip = s.getsockname()[0]
            s.close()
            return local_ip
        except Exception:
            return '127.0.0.1'
    try: ip = get('https://api.ipify.org').content.decode('utf8')
    except: ip = get_local_ip()
    print('My public IP address is: {}'.format(ip))
    return ip

@app.route('/', methods=['GET', 'POST'])
def home():
    """
    Handles HTTP requests to the root URL ('/').

    This function is responsible for handling both GET and POST requests. It checks for user sessions, 
    handles search queries, and returns relevant table data in JSON format or renders a blank HTML template if no user is logged in.

    Parameters:
    None

    Returns:
    A JSON response containing search results or a rendered HTML template.
    """
    # print(request.remote_addr,get_my_ip())
    # if session.get('user', None):
    #     # session['user']['userinfo']['admin'] = isUserAllowed(session['user']['userinfo']['email'])
    #     # if session['user'] == 'local':
    #     #     session['user']['userinfo']['admin'] = True
    #     if isUserAllowed(session['user']['userinfo']['email']):
    #         session['user']['userinfo']['admin'] = True
    if session.get('user') and isinstance(session['user'], dict) and 'userinfo' in session['user']:
        # session['user']['userinfo']['admin'] = True
        # else:
        #     session['user']['userinfo']['admin'] = False
        table_data = None # doing this so that it does not get referenced before assginment
        book = request.args.get('book', None)
        if request.method == 'POST':
            data = request.get_json(silent=True)
            if data:
                query = data['query']
                attribute = data['attribute']
                book = data['book']
                # print(query)
            
            if attribute == 'Full_Text':
                song_order = []
                table_data = []
                links = json.loads(songSearch(query)) # returns a list of links ex: <a class="list-group-item list-group-item-action" href="/song/New/300">300: Օրհնյալ Սուրբ Հոգի, մեծ Մխիթարիչ,</a>
                from re import findall
                for link in links:
                    song = {}
                    book = findall('/song/(.*?)/', link)[0]
                    try:
                        song_num = findall(r'\d+', link)[0]
                    except IndexError:
                        song_num = None
                    # table_data[book] = {new:[1,3],old:[2]} # correct order is 3,1,2
                    if song_num: # edge case where I pick up songNum from index
                        song[song_num] = getSong(book, song_num)
                        song[song_num]["book"] = book
                        lyrics = song_lyrics[book][song_num][:100]
                        clean_lyrics = re.sub("   ",'',(re.sub(r'[:,.(0-9)\n]+','',lyrics)))
                        song[song_num]['lyrics'] = clean_lyrics#[:100]
                        song_order.append(song_num)
                        table_data.append(song)
                        
                return json.dumps(table_data, ensure_ascii=False)
            
            if book and attribute and not query:
                # print(load_table_data(book=book))
                return json.dumps([load_table_data(book=book)])
            
            if query and book and attribute: # if attr is full_text
                table_data = load_table_data(book=book)
                filtered_data = {}
                query:str
                # Old code for when I had multiple search options, now just does title, as its the most often used
                # if attribute == 'SongNum': 
                #     filtered_data[query] = table_data[query]
                # else:
                    # Handle numeric queries separately
                cleaned_numeric_query = re.sub(r'[-՛:։,.\\n\s]+', ' ', query, re.MULTILINE)
                if cleaned_numeric_query.strip().isdigit():
                    filtered_data[cleaned_numeric_query] = table_data.get(cleaned_numeric_query, None)
                else:
                    query_lower = query.lower()
                    cleaned_query = re.sub(r'[^ա-ֆԱ-Ֆ\s]','',query_lower)
                    # print(cleaned_query)
                    # Build the regex string used for search
                    regex_str = fr"{cleaned_query}*[ա-ֆԱ-Ֆ].*"
                    for song_num, attrs in table_data.items():
                        # attrs is the dict containing all attrs
                        if re.match(regex_str, table_data[song_num]["Title"].lower()):
                            match = table_data[song_num]["Title"]
                            # print(f"Found a potential match at: { match }")
                            filtered_data[song_num] = attrs
                        # # Search only in the specified attribute
                        # found_song_title = str(attr[attribute]).lower()
                        # print(f"Found in dict: {found_song_title}")
                        # found_song_title = re.sub(r'[^\w\s]','',found_song_title)
                        # # found_song_title = re.sub(r'[\\n\s]', ' ', found_song_title)
                        # print(f"Found in dict(modified): {found_song_title}")
                        # print(f"Usr Input: {query_cleaned}")
                        # if found_song_title in query_cleaned:
                        #     filtered_data[song_num] = attr
                        # break
                table_data: list[dict] = [filtered_data]
                return json.dumps(table_data)
            
            elif not book:
                flash('No book selected','warning') #practically not needed anymore
        return render_template('index.html', table_data = table_data, book=book) #returns book, for continuity purposes
    else:
        # session["user"] = random() * 10 ** 17
        session["user"] = {
            "userinfo":{
                "name": "Guest",
                "email": "guest@example.com",
                "admin": False,
            },
        }
    # elif request.remote_addr == get_my_ip():
    #     session['user'] = 'local'
    #     return redirect('/')
    return render_template('index.html')

@app.route('/editsongs', methods=['GET', 'POST'])
def edit_songs():
    """
    Handles HTTP requests to the '/editsongs' route, allowing users to edit song information.
    
    If the request method is 'POST', it retrieves the song number and book from the request form,
    updates the song information based on the user's input, and saves the changes to a JSON file.
    
    The function returns a rendered template of the 'edit_songs.html' page, passing the updated song information and current values as parameters.
    
    Parameters:
    None
    
    Returns:
    A rendered template of the 'edit_songs.html' page
    """
    song_info = None
    current_values = None
    is_found = False
    if request.method == 'POST':
        
        song_num = request.form.get('songNum')
        book = request.form.get('book')
        # Bool_Lyrics = request.form.get('lyrics', False)
        
        # if Bool_Lyrics:
        #     return render_template('song.html', lyrics = openWord(song_num, book), book=book) #sending the book var inorder for the back button to function properly
        
        with open(f'{book}.json',encoding='utf-8') as f:
            data = json.load(f)
        
        songs = data.get('SongNum')  # Get the songs under the "SongNum" key
        song_info = songs.get(song_num)
        
        if song_info:
            # Save the current values before they are edited
            current_values = song_info.copy()
            if not current_values:
                flash("That song does not exist",'error')
        # else:
        #     flash("That song does not exist")
            
        if request.form.get('edit'):
            song_info['key'] = request.form.get('key')
            song_info['speed'] = request.form.get('speed')
            song_info['style'] = request.form.get('style')
            song_info['song_type'] = request.form.get('Song Type')
            # song_info['Worship_Song'] = request.form.get('Worship_Song')
            song_info['timeSig'] = request.form.get('Time Signature')
            song_info['Comments'] = request.form.get('Comments')
            songs["SongNum"] = song_info
            # with open(f'{book}.json', 'w', encoding='utf-8') as f:  # Save the changes to the same file
            #     json.dump(songs, f, indent=4, ensure_ascii=False)  # Write the whole data back to the file
        
        if request.form.get('submit'):
            song_info['key'] = request.form.get('key')
            song_info['speed'] = request.form.get('speed')
            song_info['style'] = request.form.get('style')
            song_info['song_type'] = request.form.get('Song Type')
            # song_info['Worship_Song'] = request.form.get('Worship_Song')
            song_info['timeSig'] = request.form.get('Time Signature')
            song_info['Comments'] = request.form.get('Comments')
            songs[song_num] = song_info
            # if song_info:
            #     return flash("That song does not exist",'error')
        with open(f'{book}.json', 'w', encoding='utf-8') as f:  # Save the changes to the same file
                json.dump(data, f, indent=4, ensure_ascii=False)  # Write the whole data back to the file
            
    return render_template('edit_songs.html', song_info=song_info, current_values=current_values) #add pretty=json.dumps(session.get('user'), indent=4) for debuging auth

    # return render_template('edit_songs.html')

@app.route('/tsank', methods=['GET','POST'])
def tsank():
    """
    Handles HTTP requests to the '/tsank' route, supporting both GET and POST methods.
    
    Retrieves the 'temmas' value from the request form data and uses it to load a JSON file.
    
    If 'temmas' is not None, it extracts a number from the value, uses it to index into the loaded JSON data,
    and renders the 'temas.html' template with the selected data.
    
    If 'temmas' is None, it renders the 'tema.html' template with the original 'temmas' value and an empty list.
    
    Returns a rendered HTML template.
    """
    temma = request.form.get("temmas", None)
    temmalist = None
    if temma:#checks to make sure it is not none
        from re import findall
        temmaNumber = findall(r"\d+", temma)
        with open("templates/temmas_new.json", 'r', encoding='utf-8') as f:
            temmalist = json.load(f)
        temmalist = temmalist[int(temmaNumber[0])-1]
        return render_template('temas.html', temmalist=temmalist)
            
    return render_template("tema.html", temmas=temma,temmalist=temmalist)

@app.route('/tsank_a_z', methods=['GET','POST'])
def tsank_A_Z():
    """
    Handles HTTP requests to the '/tsank_a_z' endpoint, accepting both GET and POST methods.
    Returns the rendered 'tsank_A_Z.html' template.
    """
    return render_template('tsank_A_Z.html')

@app.route('/tsank_a_z/<book>/<letter>', methods=['GET','POST'])
def tsank_letter(book, letter):
    # lookup_table = ['Ա','Բ','Գ','Դ','Ե','Զ','Է','Ը','Թ','Ժ','Ի','Լ','Խ','Ծ','Կ','Հ','Ձ','Ղ','Ճ','Մ','Յ','Ն','Շ','Ո','Չ','Պ','Ջ','Ս','Վ','Տ','Ց','Ու','Փ','Ք','Եւ','Օ']
    # lookup_index = lookup_table.index(letter)
    with open('first_letter.json', 'r', encoding='utf-8') as f:
        first_letter = json.load(f)
    selected_letter = [first_letter[book][letter][songnum] for songnum in first_letter[book][letter]]
    # selected_letter = []
    # for songnum in first_letter[book][letter]:
    #     selected_letter.append(first_letter[book][letter][songnum])
    return render_template('tsank_letter.html', selected_letter = selected_letter)

@app.route('/past_songs', methods=['GET','POST'])
def check_past_songs():
    """
    Handles the '/past_songs' route, accepting both GET and POST requests.
    
    When a POST request is made, it retrieves the 'SongNumTuple' from the request form,
    parses it to extract the song number and book, and then uses these values to search
    for past songs. If past songs are found, it formats the song titles and redirects the
    user to the 'display_song' route. If no past songs are found, it renders the 
    'check_past_songs.html' template.
    
    Parameters:
    None
    
    Returns:
    A redirect response to the 'display_song' route or a rendered 'check_past_songs.html' template.
    """
    from scanningDir import songSearch
    SongNumTuple = request.form.get('SongNumTuple')
    if SongNumTuple:
        pair = SongNumTuple.split(', ') # Song Number: ${element.songNum}, Book: ${element.book}
        SongNum = pair[0].split(': ')[1]
        book = pair[1].split(': ')[1]
        with open('wordSongsIndex.json', 'r', encoding='utf-8') as f:
            wordSongsIndex = json.load(f)
    
        with open('REDergaran.json', 'r', encoding='utf-8') as f:
            REDergaran = json.load(f)
        
        past_songs = songSearch(SongNum, book)
        if past_songs != None:
            for songs in past_songs:
                song_titles = []
                for song_pair in songs['songs']:
                    if not (None in song_pair):
                        # each song is a tuple ie: ('Old', '495')
                        # here I am simply adding a tuple(list(title))
                        title = ""
                        if song_pair[0] == 'Old':
                            if song_pair[1] in wordSongsIndex['SongNum']:
                                title = wordSongsIndex['SongNum'][song_pair[1]]["Title"]
                                title = title.split('\n')[0]
                        else:
                            if song_pair[1] in REDergaran['SongNum']:
                                title = REDergaran['SongNum'][song_pair[1]]["Title"]#REDergaran.get(['SongNum'][song_pair[1]]["Title"],None) # doing this to try to account for unusual song nums such as '32121'
                                title = title.split('\n')[0]
                        song_titles.append(f'''<a class="list-group-item list-group-item-action" href="{url_for('display_song',book=song_pair[0],songnum=song_pair[1])}">{song_pair[1]}: {title}</a>''')
                songs['songs'] = song_titles
        # return render_template('song.html', lyrics = openWord(SongNum, book), past_songs = past_songs)
        return redirect(url_for('display_song'), book=book,songnum=SongNum)#, title=title) #sending the book var inorder for the back button to function properly
    return render_template('check_past_songs.html')

#helper function for check_past_songs(), used with fetch api
@app.route('/past_songs/<songnum>', methods=['GET','POST'])
def past_songs(songnum):
    """
    Handles HTTP requests to the '/past_songs/<songnum>' route. 
    This function takes a song number and book as input, 
    searches for past songs matching the given song number and book, 
    and returns a rendered template with the past songs data.

    Parameters:
        songnum (str): The song number to search for past songs.

    Returns:
        A rendered HTML template with the past songs data.
    """
    past_songs = None
    SongNum = songnum
    data = request.get_json()
    book = data['book']
    # load the song database
    with open('wordSongsIndex.json', 'r', encoding='utf-8') as f:
        wordSongsIndex = json.load(f)
    with open('REDergaran.json', 'r', encoding='utf-8') as f:
        REDergaran = json.load(f)
    
    from scanningDir import songSearch
    past_songs = songSearch(SongNum, book)
    if past_songs != None:
        for songs in past_songs:
            song_titles = []
            for song_pair in songs['songs']:
                if not (None in song_pair):
                    title = ""
                    if song_pair[0] == 'Old':
                        if song_pair[1] in wordSongsIndex['SongNum']:
                            title = wordSongsIndex['SongNum'][song_pair[1]]["Title"]
                            title = title.split('\n')[0]
                    else:
                        if song_pair[1] in REDergaran['SongNum']:
                            title = REDergaran['SongNum'][song_pair[1]]["Title"]#REDergaran.get(['SongNum'][song_pair[1]]["Title"],None) # doing this to try to account for unusual song nums such as '32121'
                            title = title.split('\n')[0]
                    song_titles.append(f'''<a class="list-group-item list-group-item-action" href="{url_for('display_song',book=song_pair[0],songnum=song_pair[1])}">{song_pair[1]}: {title}</a>''')
            songs['songs'] = song_titles
            
    return render_template('pastsongtemplate.html', past_songs=past_songs)

@app.route('/newSundaySong', methods=['GET', 'POST'])
def newSundaySong():
    if request.method == 'POST':
        data = request.get_json()
        only_first_two_songs = data['only_first_two_songs']
        only_worship_songs = data['only_worship_songs']
        only_last_two_songs = data['only_last_two_songs']
        from song_curator import find_sunday_song
        reccomened_songs = find_sunday_song(only_first_two_songs, only_worship_songs, only_last_two_songs)
        
        return jsonify(reccomened_songs) if reccomened_songs else jsonify(None)
    return render_template('newSundaySongs.html')

@app.route('/weekdaySong', methods=['GET', 'POST'])
def weekdaySong():
    if request.method == 'POST':
        from song_curator import get_weekday_song
        reccomened_songs = get_weekday_song()
        return jsonify(reccomened_songs)
    return render_template('newWeekdaySongs.html')

@app.route('/transliterate', methods=['GET'])
def transliterate():
    session['transliterate'] = True
    return render_template('transliterate.html')

@app.route('/song/<book>/<songnum>/lyrics', methods=['GET','POST'])
def get_song_lyrics(book,songnum):
    if request.method == 'POST':
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor() as exec:
            future = exec.submit(openWord,songnum,book)
            lyrics = future.result()
        return jsonify(lyrics)
    return jsonify(None)

@app.route('/song/<book>/<songnum>/altsong', methods=['GET','POST'])
def posible_alt_song(songnum,book): # COuld also do only num,book,lyrics
    book = 'old' if book == 'wordsongsindex' or book == 'old' else 'new' #book == 'old' might seem redundanct, but considering songs can also be under both old and wordsongindex, it's better to be safe
    request_data = request.get_json()
    lyrics = request_data['lyrics']
    try:
        results = search_engine.search(lyrics)[:2] # get first two
        # print(results)
        if results[1][2] > 0.45: # check probability
            if results[1][0] != book: #and results[1][1] != songnum: # If they are of the same book and num, NO return
                # print(results[1])
                return jsonify(results[1])
            elif results[1][1] != songnum:
                # print(results[1])
                return jsonify(results[1])
            else:
                # print(results[0])
                return jsonify(results[0])
        else:
            return jsonify(None)
    except:
        return jsonify(None)
    

@app.route('/known_songs', methods=['GET','POST'])
def known_songs():# add some func to be able to go backwards
    from known_songs import update_known_songs, get_skipped_songs
    if request.method == "POST":
        request_data = request.get_json()
        # these are neccesary either way, just don't want to set unneccesary vars
        skipped = request_data['skipped']
        songnum = request_data['songId']
        book = request_data['book']
        if not skipped:
            isHoliday = request_data['isHoliday']
            isSunday = request_data['isSunday']
            isWeekday = request_data['isWeekday']
            known = request_data['choirKnows']
            
            # print(songnum)
            # print(book)
            # print(isHoliday)
            # print(isSunday)
            # print(isWeekday)
            # print(known)
            update_known_songs(book, songnum, isHoliday=isHoliday, isSunday=isSunday, isWeekday=isWeekday, known=known)
        else:
            # print(songnum)
            # print(book)
            # print("Skipped")
            update_known_songs(book, songnum, skipped=skipped)
        
    return render_template('known_songs.html', skipped_songs = get_skipped_songs())

@app.route('/known_songs/newSong', methods=['GET'])
def get_unknown_songs():# add some func to be able to go backwards
    from known_songs import get_a_song
    song, book = get_a_song()
    import concurrent.futures
    with concurrent.futures.ThreadPoolExecutor() as exec:
        future = exec.submit(openWord,song,book)
        lyrics = future.result()
    return jsonify([lyrics, book, song])

@app.route('/known_songs/getSong', methods=['GET'])
def get_skipped_songs():# add some func to be able to go backwards
    book = request.args.get("book")
    song = request.args.get("songnum")
    import concurrent.futures
    with concurrent.futures.ThreadPoolExecutor() as exec:
        future = exec.submit(openWord,song,book)
        lyrics = future.result()
    return jsonify([lyrics, book, song])

@app.route('/song_matcher')
def song_matcher():
    
    return render_template('song_matching.html',song1='',song2='')

# @app.route('/song_analysis', methods=['GET', 'POST'])
def song_analysis():
    """
    Analyzes song data to provide insights on service patterns and song usage.
    Returns analysis through the song_analysis.html template.
    """
    # Load your song data
    with open('REDergaran.json', 'r', encoding='utf-8') as f:
        red_songs = json.load(f)
    with open('wordSongsIndex.json', 'r', encoding='utf-8') as f:
        word_songs = json.load(f)
    
    # Initialize counters
    key_distribution = {}
    tempo_distribution = {}
    style_distribution = {}
    type_distribution = {}
    total_songs = 0
    
    def analyze_tempo_ranges(tempo_distribution):
        ranges = {
            'Slow (<=72)': 0,
            'Medium (73-108)': 0,
            'Fast (109-144)': 0,
            'Very Fast (>144)': 0
        }
        
        for tempo, count in tempo_distribution.items():
            try:
                tempo_val = float(tempo)
                if tempo_val <= 72:
                    ranges['Slow (<=72)'] += count
                elif tempo_val <= 108:
                    ranges['Medium (73-108)'] += count
                elif tempo_val <= 144:
                    ranges['Fast (109-144)'] += count
                else:
                    ranges['Very Fast (>144)'] += count
            except ValueError:
                continue
        return ranges

    def analyze_key_relationships(key_distribution):
        major_keys = {}
        minor_keys = {}
        
        for key, count in key_distribution.items():
            if 'm' in key.lower():
                minor_keys[key] = count
            else:
                major_keys[key] = count
                
        return {
            'major': major_keys,
            'minor': minor_keys,
            'major_total': sum(major_keys.values()),
            'minor_total': sum(minor_keys.values())
        }

    def analyze_combinations(songs_data):
        combinations = {}
        for song_data in songs_data['SongNum'].values():
            key = song_data.get('key', '').strip()
            style = song_data.get('style', '').strip()
            if key and style:
                combo = f"{key} - {style}"
                combinations[combo] = combinations.get(combo, 0) + 1
        
        # Get top 10 most common combinations
        return dict(sorted(combinations.items(), key=lambda x: x[1], reverse=True)[:10])

    def process_songs(songs_data):
        nonlocal total_songs
        for song_num, song_data in songs_data['SongNum'].items():
            total_songs += 1
            key = song_data.get('key', '').strip()
            speed = song_data.get('speed', '').strip()
            style = song_data.get('style', '').strip()
            song_type = song_data.get('song_type', '').strip()
            
            # Only count non-empty values
            if key and key.lower() != 'unknown':
                key_distribution[key] = key_distribution.get(key, 0) + 1
            if speed and speed.lower() != 'unknown':
                tempo_distribution[speed] = tempo_distribution.get(speed, 0) + 1
            if style and style.lower() != 'unknown':
                style_distribution[style] = style_distribution.get(style, 0) + 1
            if song_type and song_type.lower() != 'unknown':
                type_distribution[song_type] = type_distribution.get(song_type, 0) + 1
    
    # Process both song collections
    process_songs(red_songs)
    process_songs(word_songs)

    # Sort tempo distribution by speed if possible
    try:
        tempo_distribution = dict(sorted(tempo_distribution.items(), 
                                       key=lambda x: float(x[0])))
    except ValueError:
        # If conversion fails, keep original order
        pass

    # Get additional analyses
    tempo_ranges = analyze_tempo_ranges(tempo_distribution)
    key_relationships = analyze_key_relationships(key_distribution)
    combinations = analyze_combinations(red_songs) # You might want to merge both song collections here
    
    analysis_data = {
        'key_distribution': key_distribution,
        'tempo_distribution': tempo_distribution,
        'style_distribution': style_distribution,
        'type_distribution': type_distribution,
        'total_songs': total_songs,
        'songs_with_attributes': {
            'keys': len(key_distribution),
            'styles': len(style_distribution),
            'types': len(type_distribution),
            'tempos': len(tempo_distribution)
        },
        'tempo_ranges': tempo_ranges,
        'key_relationships': key_relationships,
        'combinations': combinations
    }
    
    return render_template('song_analysis.html', analysis=analysis_data)



if __name__ == '__main__':
    print("Barev Dzez, ev bari galust MOSO-i system....\nLaunching Server...")
    app.run(debug=True, host='0.0.0.0', port=env.get("PORT", 5002 if os.name == 'posix' else 5000), 
            #ssl_context=(r'C:\Certbot\live\songinfo.us.to\fullchain.pem',r'C:\Certbot\live\songinfo.us.to\privkey.pem')
            )
    # try: app.run(debug=True, host='0.0.0.0', port=env.get("PORT", 5000))
    # except: app.run(debug=True, host='0.0.0.0', port=env.get("PORT", 5001))
    