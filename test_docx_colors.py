import os
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.shared import RGBColor

import app as app_module


class DocxColorRenderingTests(unittest.TestCase):
    def test_colored_docx_is_rendered_safely_and_accurately(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / 'colored-test.docx'
            cache_dir = temp_path / 'htmlsongs'

            doc = Document()
            doc.add_paragraph('[start:song]')

            number = doc.add_paragraph()
            number_run = number.add_run('101')
            number_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

            lyric = doc.add_paragraph()
            lyric.add_run('Plain ')
            blue_run = lyric.add_run('blue')
            blue_run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
            lyric.add_run(' and ')
            green_run = lyric.add_run('green')
            green_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            lyric.add_run().add_break()
            lyric.add_run('next line')

            doc.add_paragraph("<script>alert('not HTML')</script>")
            doc.add_paragraph('')
            doc.add_paragraph('[end:song]')
            doc.save(source_path)

            self.assertTrue(source_path.is_file())
            song_nums = app_module.saveHtml(str(source_path), 'colored-test.docx', str(cache_dir))
            rendered_docx = (cache_dir / 'colored-test.docx.txt').read_text(encoding='utf-8')

            self.assertEqual(song_nums, ['101'])
            self.assertIn('<div id="song-1">', rendered_docx)
            self.assertIn('<span style="color: #FF0000;">101</span>', rendered_docx)
            self.assertIn('<span style="color: #0000FF;">blue</span>', rendered_docx)
            self.assertIn('<span style="color: #008000;">green</span>', rendered_docx)
            self.assertIn('<br>next line', rendered_docx)
            self.assertIn('&lt;script&gt;alert(&#x27;not HTML&#x27;)&lt;/script&gt;', rendered_docx)
            self.assertNotIn("<script>alert('not HTML')</script>", rendered_docx)

            original_cwd = os.getcwd()
            try:
                os.chdir(temp_path)
                with patch.object(app_module, 'all_past_songs', {
                    'colored-test.docx': {'songList': "[('New', '101')]"},
                }):
                    response = app_module.app.test_client().get('/docx/colored-test.docx')
            finally:
                os.chdir(original_cwd)

            self.assertEqual(response.status_code, 200)
            self.assertIn(b'<span style="color: #0000FF;">blue</span>', response.data)
            self.assertIn(b'<span style="color: #008000;">green</span>', response.data)


class SongDisplayTests(unittest.TestCase):
    def test_unknown_song_returns_not_found(self):
        response = app_module.app.test_client().get('/song/old/not-a-real-song-number')

        self.assertEqual(response.status_code, 404)


if __name__ == '__main__':
    unittest.main()
