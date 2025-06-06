#This file/module serves as a helper for my newGui.pyw app, otherwise known as MOSO
import os, docx, time, re
import json
from docx.shared import Pt,RGBColor



def getTemplate():
    wordFile = open("Choir Songs Template.dotx", 'r', encoding='utf_8')
    template = wordFile.read()
    wordFile.close
    return template

month = time.strftime('%m')
year = time.strftime('%y')
fullYear = time.strftime('%Y')
day = time.strftime('%d')


def process(filePath):
    text = ""
    doc = docx.Document(filePath)
    for p in doc.paragraphs:
        text += p.text + "\n"
    return text

def FindNum(song):
    # print(song)
    return song['song'][0]['text']

def getDocTextAndIndentation(filePath:str, my_doc):
    """Reads a DOCX file and returns a docx file with the text"""
    user = os.environ.get("USERNAME")
    doc = docx.Document("C:/Users/" + user + "/OneDrive/" + filePath)
    first = True # to only run on the first line caught, on a song by song basis
    for p in doc.paragraphs:
        first_line_indent = p.paragraph_format.first_line_indent
        left_indent = p.paragraph_format.left_indent
        right_indent = p.paragraph_format.right_indent
        Placeholder = my_doc.add_paragraph().clear()
        for aWord in p.text:
            if re.match(r"\d+", aWord) or aWord == "(" or aWord == ")":
                run = Placeholder.add_run(aWord)
                run.font.color.rgb = RGBColor(255, 0, 0) # Color for red
            else:
                run = Placeholder.add_run(aWord)
                run.font.color.rgb = RGBColor(0, 0, 0) # Color for white
       
        Placeholder.paragraph_format.space_after = 0
        if first_line_indent is not None:
            Placeholder.paragraph_format.first_line_indent = first_line_indent
        if left_indent is not None:
            Placeholder.paragraph_format.left_indent = left_indent
        if right_indent is not None:
            Placeholder.paragraph_format.right_indent = right_indent
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(22)
        
    return my_doc

def getRandomDoc():
    from random import randint
    from glob import glob
    rDoc = randint(0, 3) #Choose a random number from 4 numbers
    DrivePath = os.environ.get("OneDrive")
    posible_rand_docs = glob(DrivePath+"\\*.docx")
    return docx.Document(posible_rand_docs[rDoc])

#parses the data inputed and sends back a python-docx file object
def getPcSongs(songs, imp, user):
    """parses the data inputed and sends back a python-docx file object

    Args:
        songs (list): A list containing all of the song numbers requested
        imp (list): used to denote either 'n'ew or 'o'ld databases
        user (str): a str denoting what pc this is being run on

    Raises:
        Wrong file name: thrown by a checker that checks to see if the filePath matchs with what is found

    Returns:
        docx: a word file containing the requested songs
    """
    my_doc = getRandomDoc()
    # DrivePath = os.environ.get("OneDrive")
    # my_doc = docx.Document(F"{DrivePath}\\Choir Songs Template - Alt - Copy.docx")
    
    for x in songs:
        x = str(x)
        y = songs.index(x)
        if 'n' in imp[y]:
            
        
            with open("REDergaran.json", 'r', encoding='utf-8') as f:
                ergaran = json.load(f)
            filePath = ""
            if x in ergaran["SongNum"]:
                filePath = ergaran["SongNum"][x]["latestVersion"]
                
                Placeholder = my_doc.add_paragraph()
                run = Placeholder.add_run("[start:song]")
                run.font.color.rgb = RGBColor(127, 165, 249)
                
                getDocTextAndIndentation(filePath=filePath, my_doc=my_doc)
                
                Placeholder = my_doc.add_paragraph()
                run = Placeholder.add_run("[end:song]")
                run.font.color.rgb = RGBColor(127, 165, 249)
            else:
                Placeholder = my_doc.add_paragraph()
                run = Placeholder.add_run("[start:song]")
                run.font.color.rgb = RGBColor(127, 165, 249)
                getDocTextAndIndentation(filePath="RED Words/" + x + ".docx", my_doc=my_doc)
                Placeholder = my_doc.add_paragraph()
                run = Placeholder.add_run("[end:song]")
                run.font.color.rgb = RGBColor(127, 165, 249)
                
        else:
            #Get file path from 'old' database
            with open("wordSongsIndex.json", 'r', encoding='utf-8') as f:
                OldErgaran = json.load(f)
            filePath = ""
            if x in OldErgaran["SongNum"]:
                filePath = OldErgaran["SongNum"][x]["latestVersion"]
            else:
                my_doc.add_paragraph("Error: FileNotFoundError \nSong: " + x + " Old, Could not be located ")
            
            Placeholder = my_doc.add_paragraph()
            run = Placeholder.add_run("[start:song:old]")
            run.font.color.rgb = RGBColor(127, 165, 249)
            getDocTextAndIndentation(filePath=filePath, my_doc=my_doc)
            Placeholder = my_doc.add_paragraph()
            run = Placeholder.add_run("[end:song:old]")
            run.font.color.rgb = RGBColor(127, 165, 249)


    
    my_doc.add_page_break()
    
    return my_doc

def getPosibleSongs(songs, imp):
    posSongList = []
    user = os.environ.get("USERNAME")
    for x in songs:
        x = str(x)
        y = songs.index(x)
        if 'n' in imp[y]:
            
            try:
                with open("REDergaran.json", 'r', encoding='utf-8') as f:
                    ergaran = json.load(f)
                filePath = ""
                if x in ergaran["SongNum"]:
                    filePath = ergaran["SongNum"][x]["latestVersion"]
                    filePath = re.sub('.*/+',"",filePath)
                posSongList.append(filePath)
            except:
                posSongList.append("RED Words/" + str(x))
        else:
            
            try:
                with open("wordSongsIndex.json", 'r', encoding='utf-8') as f:
                    OldErgaran = json.load(f)
                filePath = ""
                if x in OldErgaran["SongNum"]:
                    filePath = OldErgaran["SongNum"][x]["latestVersion"]
                    filePath = re.sub('.*/+',"",filePath)
                posSongList.append(filePath)
            except:
                posSongList.append("Could not find old song: {}".format(x))
        

    
    return posSongList

