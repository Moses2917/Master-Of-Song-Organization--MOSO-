import os
import re
from html import escape

from docx import Document


_RGB_HEX = re.compile(r"^[0-9A-Fa-f]{6}$")


def _render_run(run) -> str:
    """Render one Word run while preserving its direct RGB font color."""
    text = escape(run.text).replace("\r\n", "\n").replace("\r", "\n").replace("\n", "<br>")
    rgb = run.font.color.rgb
    if rgb and _RGB_HEX.fullmatch(str(rgb)):
        return f'<span style="color: #{str(rgb).upper()};">{text}</span>'
    return text


def render_docx_html(filepath='', doc=None) -> str:
    """Convert a DOCX into safe HTML while retaining paragraphs, breaks, and RGB run colors."""
    if doc is None:
        doc = Document(filepath)

    html_parts = []
    song_number = 0
    song_open = False

    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text.lower()
        starts_song = "[start:song" in paragraph_text
        ends_song = "[end:song" in paragraph_text

        if starts_song:
            if song_open:
                html_parts.append("</div>")
            song_number += 1
            html_parts.append(f'<div id="song-{song_number}">')
            song_open = True

        paragraph_html = "".join(_render_run(run) for run in paragraph.runs)
        if not paragraph_html and paragraph.text:
            paragraph_html = escape(paragraph.text).replace("\n", "<br>")
        html_parts.append(f"<p>{paragraph_html or '&nbsp;'}</p>")

        if ends_song and song_open:
            html_parts.append("</div>")
            song_open = False

    if song_open:
        html_parts.append("</div>")

    return "".join(html_parts)


def get_colored_text(filepath='', doc=None) -> dict:

    # Open the document
    if not doc:
        doc = Document(filepath)
    colored_text = {}
    # Access a specific paragraph (e.g., the first one)
    for paragraph in doc.paragraphs:
        # Iterate through runs in the paragraph to check their font colors
        for run in paragraph.runs:
            # Get the color format object
            color_format = run.font.color
            
            # if color_format.type is None:
            #     print(f'Text: "{run.text}", Color: None (default)')
            if hasattr(color_format, 'rgb') and color_format.rgb:
                # Get the RGB value
                rgb_value = color_format.rgb

                rgb_str = str(rgb_value)
                if rgb_str != "000000": # Only get colored text
                    colored_text[run.text] = rgb_str
    return colored_text


def get_all_colored_text(filepath='', doc=None) -> list[str]:
    """Returns a html code in a list of all text and includes color"""
    
    if not doc:
        doc = Document(filepath)
    
    colored_text = []
    song_nums = []
    song_counter = 0
    start_song = False
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == '':
            colored_text.append("<br>")
            continue
            
        # Process runs to preserve colors, but handle line breaks properly
        full_paragraph_html = ""
        current_line = ""
        
        for run in paragraph.runs:
            text = run.text
            
            # Handle song markers
            if start_song:
                if song_counter != 1:
                    colored_text.append('</div>')
                colored_text.append(f'<div id="song-{song_counter}">')
                song_nums.append(text.strip())
                start_song = False
                
            if "start" in text:
                start_song = True
                song_counter += 1
            
            # Split by line breaks but preserve color info
            lines = text.split('\n')
            
            for i, line in enumerate(lines):
                if i > 0:  # New line found, finish current paragraph
                    if current_line.strip():
                        colored_text.append(f"<p>{current_line}</p>")
                    current_line = ""
                
                if line.strip():  # Only process non-empty lines
                    text_color = run.font.color.rgb
                    if bool(text_color) and str(text_color) != "000000":
                        current_line += f'<span style="color: #{text_color};">{line}</span>'
                    else:
                        current_line += line
        
        # Add the final line if any
        if current_line.strip():
            colored_text.append(f"<p>{current_line}</p>")
    
    # Add closing div if needed
    if song_counter > 0:
        colored_text.append('</div>')
    
    return tuple([colored_text, song_nums])

def make_colored_doc(filepath='', doc=None):
    """Make the doc in the soecufed location, but ensures that the text color is the same as in the orignal doc.

    Args:
        filepath (str, optional): A direct path to the document. Defaults to ''.
        doc (Document, optional): A Document objext for if you wish to just pass that in. Defaults to None.
    """
    if not doc:
        doc = Document(filepath)
    else:
        doc_true = True # If you passed in a document variabe, into the function
    
    colored_text = []
    song_nums = []
    song_counter = 0
    start_song = False

if '__main__' == __name__:
    if os.name == "nt":
        filepath = r"C:\Users\moses\OneDrive\Երգեր\Պենտեկոստե/2025/Պենտեկոստե.docx"#06.2025\06.01.25.docx"
    else:
        filepath = r"/Users/movsesmovsesyan/Library/CloudStorage/OneDrive-Personal/Երգեր/Պենտեկոստե/2025/Պենտեկոստե.docx"
    # colored_text = get_all_colored_text(filepath)
    # for line in colored_text[0]:
    #     print(line)

    # print(colored_text[1])
